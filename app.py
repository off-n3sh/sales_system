"""
DREAMLAND ZERO - DISTRIBUTED FLASK APPLICATION

"""
# ════════════════════════════════════════════════════════════
# IMPORTS
# ════════════════════════════════════════════════════════════
from flask import Flask, request, jsonify, session, redirect, url_for, render_template, flash
from functools import wraps
from pymongo import MongoClient
from pymongo.errors import ConnectionFailure, ServerSelectionTimeoutError, OperationFailure
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import datetime, timedelta
import logging
import re
import secrets
import pytz
import json
from typing import Union, List, Dict, Optional, Tuple
from bson import ObjectId
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from io import BytesIO
from flask import Response
from dateutil import parser
from apscheduler.schedulers.background import BackgroundScheduler
import os
import sys
import platform
from werkzeug.utils import secure_filename

# ════════════════════════════════════════════════════════════
# FLASK & LOGGING SETUP
# ════════════════════════════════════════════════════════════
app = Flask(__name__)
app.config['SECRET_KEY'] = secrets.token_hex(32)

NAIROBI_TZ = pytz.timezone('Africa/Nairobi')

os.makedirs('logs', exist_ok=True)

# ════════════════════════════════════════════════════════════
# LOGGING SETUP
# ════════════════════════════════════════════════════════════
class CustomFormatter(logging.Formatter):
    icons = {
        logging.INFO: "ℹ️ ",
        logging.WARNING: "⚠️ ",
        logging.ERROR: "❌",
        logging.CRITICAL: "🚨"
    }

    def format(self, record):
        icon = self.icons.get(record.levelno, "  ")
        if not record.msg.startswith(icon):
            record.msg = f"{icon} {record.msg}"
        return super().format(record)

log_format = '%(asctime)s | %(levelname)-8s | %(name)-10s | %(message)s'
formatter = CustomFormatter(log_format, datefmt='%H:%M:%S')

# Handlers
main_handler = logging.FileHandler('logs/app.log')
stream_handler = logging.StreamHandler()
backup_handler = logging.FileHandler('logs/backup.log')

for h in [main_handler, stream_handler, backup_handler]:
    h.setFormatter(formatter)

log = logging.getLogger('app')
logger = log
log.setLevel(logging.INFO)
log.addHandler(main_handler)
log.addHandler(stream_handler)

backup_log = logging.getLogger('backup')
backup_log.addHandler(backup_handler)

# ════════════════════════════════════════════════════════════
# CONFIGURATION
# ════════════════════════════════════════════════════════════
class Config:
    MONGO_USER = os.getenv('MONGO_USER')
    MONGO_PASS = os.getenv('MONGO_PASS')
    MONGO_HOST = os.getenv('MONGO_HOST', 'localhost')
    
    # Peer Detection
    MY_HOST = platform.node()
    PEER_HOST = "dream1" if "admin1" in MY_HOST else "admin1"
    
    DB_NAME = "dreamland_zero"
    CONNECTION_TIMEOUT = 3000
    
    BACKUP_TIME = {"hour": 20, "minute": 0}
    HEALTH_CHECK_INTERVAL = 30
    STOCK_CACHE_TIMEOUT = timedelta(hours=1)

    @classmethod
    def get_uri(cls, host):
        auth = f"{cls.MONGO_USER}:{cls.MONGO_PASS}@" if cls.MONGO_USER else ""
        return f"mongodb://{auth}{host}:27017/{cls.DB_NAME}?authSource=dreamland_zero&serverSelectionTimeoutMS={cls.CONNECTION_TIMEOUT}"

# ════════════════════════════════════════════════════════════
# DATABASE CONNECTION MANAGER
# ════════════════════════════════════════════════════════════
class MongoConnectionManager:
    def __init__(self):
        self.client = None
        self.current_host = Config.MY_HOST
        self.is_primary = True 
        self.peer_online = False
    
    def get_local_pid(self):
        try:
            if platform.system() == "Windows":
                cmd = ["tasklist", "/FI", "IMAGENAME eq mongod.exe", "/FO", "CSV", "/NH"]
                output = subprocess.check_output(cmd).decode()
                return output.split(",")[1].strip('"')
            else:
                return subprocess.check_output(["pgrep", "-f", "mongod"]).decode().strip().split('\n')[0]
        except: 
            return "N/A"

    def connect_to_local(self):
        try:
            self.client = MongoClient(Config.get_uri(os.getenv('MONGO_HOST', 'localhost')), connect=True)
            self.client.admin.command('ping')
            return self.client, "localhost", True
        except Exception as e:
            log.critical(f"💥 DATABASE OFFLINE: {e}")
            sys.exit(1)

    def check_peer_mismatch(self):
        """Background check for orders/stock count differences"""
        try:
            peer_uri = Config.get_uri(Config.PEER_HOST)
            with MongoClient(peer_uri, serverSelectionTimeoutMS=2000) as p_cli:
                p_db = p_cli[Config.DB_NAME]
                p_orders = p_db['orders'].count_documents({})
                l_orders = self.client[Config.DB_NAME]['orders'].count_documents({})
                
                if p_orders != l_orders:
                    log.warning(f"⚠️ [MISMATCH] Orders count: Peer({p_orders}) vs Local({l_orders})")
                self.peer_online = True
        except:
            self.peer_online = False

    def reconnect(self):
        if self.client: 
            self.client.close()
        self.client, self.current_host, self.is_primary = self.connect_to_local()
        initialize_collections()

# Initialize DB Manager
db_manager = MongoConnectionManager()

# ════════════════════════════════════════════════════════════
# GLOBAL COLLECTION REFERENCES
# ════════════════════════════════════════════════════════════
client = None
db = None
admin_db = None
audit_db = None
users_collection = None
orders_collection = None
stock_collection = None
metadata_collection = None
notifications_collection = None
clients_collection = None
stock_logs_collection = None
user_action_logs = None

def initialize_collections():
    """Initialize all collection references after connection"""
    global client, db, admin_db, audit_db
    global users_collection, orders_collection, stock_collection
    global metadata_collection, notifications_collection, clients_collection
    global stock_logs_collection, user_action_logs
    
    client = db_manager.client
    db = client[Config.DB_NAME]
    admin_db = client["admin"]
    audit_db = client["admin_audit"]
    
    users_collection = db["users"]
    orders_collection = db["orders"]
    stock_collection = db["stock"]
    metadata_collection = db["metadata"]
    notifications_collection = db["notifications"]
    clients_collection = db["clients"]
    stock_logs_collection = db["stock_logs"]
    user_action_logs = db["user_actions"]
    
    log.debug("✅ Collections linked")

def get_db():
    """Helper function to get the main database"""
    return db


# ════════════════════════════════════════════════════════════
# BACKUP & SCHEDULER
# ════════════════════════════════════════════════════════════
def auto_backup():
    backup_log.info("🔄 Starting local automated backup")
    try:
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        os.makedirs("backup", exist_ok=True)
        # Use existing db reference to dump data to JSON locally
        backup_data = {}
        for col_name in ["users", "orders", "stock", "metadata", "notifications", "clients"]:
            backup_data[col_name] = list(db[col_name].find())
        
        with open(f"backup/dump_{ts}.json", "w") as f:
            json.dump(backup_data, f, default=str)
        backup_log.info(f"✅ Local backup successful: dump_{ts}.json")
    except Exception as e:
        backup_log.error(f"❌ Backup failed: {e}")

def setup_schedulers():
    scheduler = BackgroundScheduler(timezone=NAIROBI_TZ)
    scheduler.add_job(auto_backup, 'cron', hour=Config.BACKUP_TIME['hour'], minute=Config.BACKUP_TIME['minute'])
    scheduler.add_job(db_manager.check_peer_mismatch, 'interval', seconds=Config.HEALTH_CHECK_INTERVAL)
    scheduler.start()
    log.info("⏰ Schedulers active")

# ════════════════════════════════════════════════════════════
# STARTUP INITIALIZATION
# ════════════════════════════════════════════════════════════
@app.before_request
def ensure_active_connection():
    if not db_manager.client:
        db_manager.reconnect()

def startup():
    process = "MAIN" if os.environ.get('WERKZEUG_RUN_MAIN') else "INIT"
    log.info("=" * 60)
    log.info(f"🚀 DREAMLAND ZERO STARTUP [{process}]")
    log.info(f"💻 PLATFORM: {platform.system()} ({platform.machine()})")
    log.info("=" * 60)
    
    if not Config.MONGO_USER or not Config.MONGO_PASS:
        log.critical("Missing MONGO_USER or MONGO_PASS environment variables!")
        sys.exit(1)

    # 1. Connect
    db_manager.client, db_manager.current_host, db_manager.is_primary = db_manager.connect_to_local()
    
    # 2. Initialize Data
    initialize_collections()
    
    # 3. Secure Indexes
    try:
        if "email_1" not in users_collection.index_information():
            users_collection.create_index("email", unique=True)
            log.info("📇 Unique index verified")
    except Exception as e:
        log.warning(f"Index creation skipped: {e}")

    # 4. App State
    app.stock_cache = {'data': None, 'timeout': Config.STOCK_CACHE_TIMEOUT}
    
    if not os.environ.get('WERKZEUG_RUN_MAIN'):
        setup_schedulers()
    
    pid = db_manager.get_local_pid()
    log.info(f"🔌 Connected to Local DB | PID: {pid}")
    log.info(f"🖥️  Node: {Config.MY_HOST} | Mode: STANDALONE")
    log.info("✅ System Ready")
    log.info("=" * 60)

# Trigger Startup
startup()

# ════════════════════════════════════════════════════════════
# HEALTH & MONITORING
# ════════════════════════════════════════════════════════════
stock_cache = {'data': None, 'version': None, 'timestamp': None, 'timeout': Config.STOCK_CACHE_TIMEOUT}
def get_server_status():
    """
    Pulls live MongoDB server state.
    Used for:
    - realtime status (today)
    - historical aggregation (yesterday, etc)
    """
    status = admin_db.command("serverStatus")

    return {
        "connections_current": status["connections"]["current"],
        "connections_total": status["connections"]["totalCreated"],
        "uptime": status["uptime"],
        "auth_success": status["security"]["authentication"]
            ["mechanisms"]["SCRAM-SHA-256"]
            ["authenticate"]["successful"]
    }

def is_mongod_running():
    """
    OS truth check.
    Mongo might accept connections even when degraded.
    """
    try:
        subprocess.check_output(["pgrep", "-a", "mongod"])
        return True
    except subprocess.CalledProcessError:
        return False



# ============================================================================
# SESSION CONFIGURATION
# ============================================================================

# Set session lifetime to 6 hours
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=6)
app.config['SESSION_REFRESH_EACH_REQUEST'] = True

# Force logout flag
FORCE_LOGOUT_FLAG_FILE = 'force_logout.lock'

def is_force_logout_active():
    """Check if admin has triggered force logout for all users"""
    return os.path.exists(FORCE_LOGOUT_FLAG_FILE)

def trigger_force_logout():
    """Admin function: Create flag to force all user to logout"""
    with open(FORCE_LOGOUT_FLAG_FILE, 'w') as f:
        f.write(str(datetime.now()))

def clear_force_logout():
    """Admin function: Remove force logout flag"""
    if os.path.exists(FORCE_LOGOUT_FLAG_FILE):
        os.remove(FORCE_LOGOUT_FLAG_FILE)
# ════════════════════════════════════════════════════════════
# AUDIT & SESSION TRACKING
# ════════════════════════════════════════════════════════════
@app.before_request
def track_user_activity():
    if request.endpoint and request.endpoint in ['static', 'auth', '/']:
        return
    
    if 'user' not in session:
        if request.path.startswith('/api/'):
            return jsonify({'status': 'unauthorized', 'reason': 'session_expired'}), 401
        return redirect(url_for('auth'))
    
    # === CHECK 1: Global Force Logout Flag (file-based) ===
    if is_force_logout_active():
        user_email = session.get('user', {}).get('email', 'unknown')
        db['session_logs'].insert_one({
            'email': user_email,
            'action': 'forced_logout',
            'reason': 'admin_triggered',
            'timestamp': datetime.now(NAIROBI_TZ),
            'ip_address': request.remote_addr,
            'user_agent': request.headers.get('User-Agent')
        })
        session.clear()
        if request.path.startswith('/api/'):
            return jsonify({'status': 'unauthorized', 'reason': 'global_logout'}), 401
        flash('Your session has been terminated. Please log in again.', 'warning')
        return redirect(url_for('auth'))

    # === CHECK 2: Per-user block/force_logout from DB ===
    user_id = session.get('user', {}).get('user_id')
    if user_id:
        db_user = db.users.find_one({'_id': ObjectId(user_id)}, {'status': 1, 'force_logout': 1})
        if db_user:
            if db_user.get('status') == 'blocked':
                user_email = session.get('user', {}).get('email', 'unknown')
                db['session_logs'].insert_one({
                    'email': user_email,
                    'action': 'forced_logout',
                    'reason': 'account_blocked',
                    'timestamp': datetime.now(NAIROBI_TZ),
                    'ip_address': request.remote_addr,
                    'user_agent': request.headers.get('User-Agent')
                })
                session.clear()
                if request.path.startswith('/api/'):
                    return jsonify({'status': 'unauthorized', 'reason': 'account_blocked'}), 401
                flash('Your account has been suspended. Contact support.', 'warning')
                return redirect(url_for('auth'))

            if db_user.get('force_logout'):
                user_email = session.get('user', {}).get('email', 'unknown')
                db['session_logs'].insert_one({
                    'email': user_email,
                    'action': 'forced_logout',
                    'reason': 'admin_force_logout',
                    'timestamp': datetime.now(NAIROBI_TZ),
                    'ip_address': request.remote_addr,
                    'user_agent': request.headers.get('User-Agent')
                })
                db.users.update_one({'_id': ObjectId(user_id)}, {'$unset': {'force_logout': ''}})
                session.clear()
                if request.path.startswith('/api/'):
                    return jsonify({'status': 'unauthorized', 'reason': 'force_logout'}), 401
                flash('You have been logged out by an administrator.', 'warning')
                return redirect(url_for('auth'))

    # === CHECK 3: Session Expiry ===
    now = datetime.now(NAIROBI_TZ)
    last_activity = session.get('last_activity')
    if last_activity:
        if isinstance(last_activity, str):
            last_activity = datetime.fromisoformat(last_activity)
        if (now - last_activity) > timedelta(hours=6):
            user_email = session.get('user', {}).get('email', 'unknown')
            db['session_logs'].insert_one({
                'email': user_email,
                'action': 'session_expired',
                'reason': 'inactivity_6_hours',
                'timestamp': now,
                'ip_address': request.remote_addr,
                'user_agent': request.headers.get('User-Agent')
            })
            session.clear()
            if request.path.startswith('/api/'):
                return jsonify({'status': 'unauthorized', 'reason': 'session_expired'}), 401
            flash('Your session expired due to inactivity. Please log in again.', 'info')
            return redirect(url_for('auth', next=request.url))

    session['last_activity'] = now.isoformat()
    session.permanent = True


# ============================================================
# HELPER FUNCTIONS
# ============================================================
UTC = pytz.utc

def to_eat(dt):
    if dt is None:
        return None
    if dt.tzinfo is None:
        return UTC.localize(dt).astimezone(NAIROBI_TZ)
    else:
        return dt.astimezone(NAIROBI_TZ)

def fetch_raw_activities(orders_collection, expenses_collection, filter_type, page, per_page):
    """Fetch activities without date filter (raw mode)"""
    activities = []
    
    if filter_type == 'all':
        # Get all orders
        orders = list(orders_collection.find().sort('date', -1))
        for order in orders:
            activities.append(format_order_simple(order))
        
        # Get all expenses
        expenses = list(expenses_collection.find().sort('date', -1))
        for expense in expenses:
            activities.append(format_expense_simple(expense))
        
        # Sort by date
        activities.sort(key=lambda x: x['date'], reverse=True)
    
    elif filter_type == 'pending':
        orders = list(orders_collection.find({'balance': {'$gt': 0}}).sort('date', -1))
        activities = [format_order_simple(o) for o in orders]
    
    elif filter_type == 'completed':
        orders = list(orders_collection.find({'balance': {'$lte': 0}}).sort('date', -1))
        activities = [format_order_simple(o) for o in orders]
    
    elif filter_type == 'expenses':
        expenses = list(expenses_collection.find().sort('date', -1))
        activities = [format_expense_simple(e) for e in expenses]
    
    elif filter_type == 'gateway':
        orders = list(orders_collection.find({
            '$or': [
                {'payment_type': 'mpesa'},
                {'payment_breakdown.mpesa': {'$exists': True}},
                {'payment_history.payment_type': 'mpesa'}
            ]
        }).sort('date', -1))
        activities = [format_order_simple(o) for o in orders]
    
    elif filter_type == 'modified':
        orders = list(orders_collection.find({'edit_tag': {'$exists': True}}).sort('date', -1))
        activities = [format_order_simple(o) for o in orders]
    
    elif filter_type == 'previous':
        # Orders opened before today but have payments today
        now = datetime.now(NAIROBI_TZ)
        today_start = now.replace(hour=0, minute=0, second=0)
        
        orders = list(orders_collection.find({
            'date': {'$lt': today_start},
            'payment_history.date': {'$gte': today_start}
        }).sort('date', -1))
        activities = [format_order_simple(o) for o in orders]
    
    # Paginate
    total = len(activities)
    total_pages = (total + per_page - 1) // per_page if total > 0 else 1
    start_idx = (page - 1) * per_page
    end_idx = start_idx + per_page
    paginated = activities[start_idx:end_idx]
    
    return {
        'data': paginated,
        'pagination': {
            'current_page': page,
            'per_page': per_page,
            'total': total,
            'total_pages': total_pages,
            'has_prev': page > 1,
            'has_next': page < total_pages
        }
    }


def format_order_simple(order):
    """Format order for raw mode (no date filter)"""
    return {
        'id': order['receipt_id'],
        'type': 'order',
        'receipt_id': order['receipt_id'],
        'user': order.get('salesperson_name', 'N/A'),
        'shop': order.get('shop_name', 'N/A'),
        'date': order['date'].isoformat(),
        'order_type': order.get('order_type', 'retail'),
        'payment': float(order.get('payment', 0)),
        'balance': float(order.get('balance', 0)),
        'total': float(order.get('payment', 0)) + float(order.get('balance', 0)),
        'status': 'completed' if float(order.get('balance', 0)) <= 0 else 'pending',
        'is_modified': 'edit_tag' in order,
        'edit_tag': order.get('edit_tag'),
        'is_previous': False
    }


def format_order_detailed(order, timestamp, date_start):
    """Format order for sorted mode"""
    return {
        'id': f"order_{order['receipt_id']}",
        'timestamp': timestamp.isoformat(),
        'type': 'order',
        'receipt_id': order['receipt_id'],
        'user': order.get('salesperson_name', 'N/A'),
        'shop': order.get('shop_name', 'N/A'),
        'open_date': order['date'].isoformat(),
        'close_date': order.get('closed_date').isoformat() if order.get('closed_date') else None,
        'order_type': order.get('order_type', 'retail'),
        'payment': float(order.get('payment', 0)),
        'balance': float(order.get('balance', 0)),
        'total': float(order.get('payment', 0)) + float(order.get('balance', 0)),
        'status': 'completed' if float(order.get('balance', 0)) <= 0 else 'pending',
        'is_modified': 'edit_tag' in order,
        'edit_tag': order.get('edit_tag'),
        'is_previous': order['date'] < date_start
    }


def format_payment_activity(order, payment, is_last, date_start):
    """Format individual payment as activity"""
    return {
        'id': f"payment_{order['receipt_id']}_{payment['date'].timestamp()}",
        'timestamp': payment['date'].isoformat(),
        'type': 'payment',
        'receipt_id': order['receipt_id'],
        'user': order.get('salesperson_name', 'N/A'),
        'shop': order.get('shop_name', 'N/A'),
        'activity_text': f"Paid KSh {payment['amount']} ({payment['payment_type'].upper()})",
        'open_date': order['date'].isoformat(),
        'close_date': order.get('closed_date').isoformat() if is_last and order.get('closed_date') else None,
        'order_type': order.get('order_type', 'retail'),
        'payment_type': payment['payment_type'],
        'amount_paid': float(payment['amount']),
        'total': float(order.get('payment', 0)) + float(order.get('balance', 0)),
        'balance': float(order.get('balance', 0)),
        'status': 'completed' if is_last and float(order.get('balance', 0)) <= 0 else 'pending',
        'is_modified': 'edit_tag' in order,
        'edit_tag': order.get('edit_tag'),
        'is_previous': order['date'] < date_start
    }


def format_order_for_date(order, date_start, date_end):
    """Format order as single activity with intelligent toggle logic"""

    # Correctly convert order date from UTC to EAT
    order_date = to_eat(order['date'])

    # Determine if this is a previous order
    is_previous = order_date < date_start

    # Get ALL payment history
    payment_history = order.get('payment_history', [])

    # Separate payments: on selected date vs before
    payments_today = []
    payments_before = []

    for p in payment_history:
        p_date = to_eat(p['date'])  # Correct UTC -> EAT conversion

        if date_start <= p_date <= date_end:
            payments_today.append({
                'time': p_date.strftime('%H:%M'),
                'date': p_date.strftime('%d/%m/%Y'),
                'payment_type': p['payment_type'],
                'amount': float(p['amount'])
            })
        elif p_date < date_start:
            payments_before.append({
                'time': p_date.strftime('%H:%M'),
                'date': p_date.strftime('%d/%m/%Y'),
                'payment_type': p['payment_type'],
                'amount': float(p['amount'])
            })

    # Group dual payments (same timestamp)
    payments_today_grouped = group_dual_payments(payments_today)
    payments_before_grouped = group_dual_payments(payments_before)

    # Calculate balances
    order_total = float(order.get('payment', 0)) + float(order.get('balance', 0))

    # Balance BEFORE selected date (for previous orders)
    balance_before_today = 0
    if is_previous:
        total_paid_before = sum(p['amount'] for p in payments_before)
        balance_before_today = order_total - total_paid_before

    # Balance AS IT WAS at end of this date — compare EAT-aware datetimes consistently
    total_paid_by_date = sum(
        p['amount'] for p in payment_history
        if to_eat(p['date']) <= date_end
    )
    balance_on_date = order_total - total_paid_by_date

    # Determine if toggle is needed
    needs_toggle = (
        len(payments_today_grouped) > 1 or
        is_previous or
        'edit_tag' in order or
        any(p.get('is_dual') for p in payments_today_grouped)
    )

    activity = {
        'id': order['receipt_id'],
        'timestamp': order_date.isoformat(),
        'type': 'order',
        'receipt_id': order['receipt_id'],
        'user': order.get('salesperson_name', 'N/A'),
        'shop': order.get('shop_name', 'N/A'),
        'open_date': order_date.isoformat(),
        'close_date': to_eat(order['closed_date']).isoformat() if order.get('closed_date') else None,
        'order_type': order.get('order_type', 'retail'),
        'payment': float(order.get('payment', 0)),
        'balance': balance_on_date,
        'total': order_total,
        'status': 'completed' if balance_on_date <= 0 else 'pending',
        'is_previous': is_previous,
        'is_modified': 'edit_tag' in order,
        'needs_toggle': needs_toggle,
        'payment_activity_today': payments_today_grouped,
        'payment_activity_before': payments_before_grouped,
        'balance_before_today': balance_before_today,
        'total_paid_today': sum(p.get('total_amount', p.get('amount', 0)) for p in payments_today_grouped),
        'edit_tag': order.get('edit_tag')
    }

    return activity


def group_dual_payments(payments):
    """Group payments by timestamp and detect dual payments"""
    if not payments:
        return []
    
    grouped = []
    i = 0
    
    while i < len(payments):
        current = payments[i]
        current_time = current['time']
        current_date = current['date']
        
        # Check if next payment has same timestamp
        if i + 1 < len(payments) and payments[i + 1]['time'] == current_time:
            # Dual payment detected
            next_payment = payments[i + 1]
            
            # Determine which is cash and which is mpesa
            if current['payment_type'] == 'cash':
                cash_amount = current['amount']
                mpesa_amount = next_payment['amount']
            else:
                cash_amount = next_payment['amount']
                mpesa_amount = current['amount']
            
            grouped.append({
                'time': current_time,
                'date': current_date,
                'payment_type': 'dual',
                'is_dual': True,
                'cash_amount': cash_amount,
                'mpesa_amount': mpesa_amount,
                'total_amount': cash_amount + mpesa_amount
            })
            
            i += 2  # Skip both payments
        else:
            # Single payment
            grouped.append({
                'time': current_time,
                'date': current_date,
                'payment_type': current['payment_type'],
                'is_dual': False,
                'amount': current['amount']
            })
            
            i += 1
    
    return grouped

def format_expense_simple(expense):
    """Format expense for raw mode"""
    user_data = expense.get('user_id', {})
    if isinstance(user_data, dict):
        user_name = f"{user_data.get('first_name', '')} {user_data.get('last_name', '')}".strip()
    else:
        user_name = str(user_data)
    
    return {
        'id': str(expense['_id']),
        'type': 'expense',
        'user': user_name or 'N/A',
        'activity_text': expense.get('description', 'No description'),
        'date': expense['date'].isoformat(),
        'category': expense.get('category', 'General'),
        'amount': float(expense.get('amount', 0)),
        'status': 'completed'
    }

def format_expense_detailed(expense):
    """Format expense for sorted mode"""
    user_data = expense.get('user_id', {})
    if isinstance(user_data, dict):
        user_name = f"{user_data.get('first_name', '')} {user_data.get('last_name', '')}".strip()
    else:
        user_name = str(user_data)

    # Correct UTC -> EAT conversion
    expense_date = to_eat(expense['date'])

    return {
        'id': f"expense_{expense['_id']}",
        'timestamp': expense_date.isoformat(),
        'type': 'expense',
        'user': user_name or 'N/A',
        'activity_text': expense.get('description', 'No description'),
        'open_date': expense_date.isoformat(),
        'close_date': expense_date.isoformat(),
        'category': expense.get('category', 'General'),
        'amount': float(expense.get('amount', 0)),
        'status': 'completed',
        'needs_toggle': False
    }


# Update fetch_sorted_activities to use the new format_order_for_date
def fetch_sorted_activities(orders_collection, expenses_collection, filter_type, date_start, date_end, page, per_page):
    """Fetch activities with date filter - single row per order with toggle logic"""
    activities = []
    
    if filter_type == 'all':
        # Get orders opened on this date OR with payments on this date
        orders = list(orders_collection.find({
            '$or': [
                {'date': {'$gte': date_start, '$lte': date_end}},
                {'payment_history.date': {'$gte': date_start, '$lte': date_end}}
            ]
        }))
        
        # Format each order
        for order in orders:
            activities.append(format_order_for_date(order, date_start, date_end))
        
        # Get expenses
        expenses = list(expenses_collection.find({
            'date': {'$gte': date_start, '$lte': date_end}
        }))
        
        for expense in expenses:
            activities.append(format_expense_detailed(expense))
    
    elif filter_type in ['pending', 'completed']:
        balance_filter = {'$gt': 0} if filter_type == 'pending' else {'$lte': 0}
        orders = list(orders_collection.find({
            'balance': balance_filter,
            '$or': [
                {'date': {'$gte': date_start, '$lte': date_end}},
                {'payment_history.date': {'$gte': date_start, '$lte': date_end}}
            ]
        }))
        
        for order in orders:
            activities.append(format_order_for_date(order, date_start, date_end))
    
    elif filter_type == 'expenses':
        expenses = list(expenses_collection.find({
            'date': {'$gte': date_start, '$lte': date_end}
        }))
        activities = [format_expense_detailed(e) for e in expenses]
    
    elif filter_type == 'gateway':
        orders = list(orders_collection.find({
            '$or': [
                {'payment_type': 'mpesa'},
                {'payment_breakdown.mpesa': {'$exists': True}},
                {'payment_history.payment_type': 'mpesa'}
            ],
            '$or': [
                {'date': {'$gte': date_start, '$lte': date_end}},
                {'payment_history.date': {'$gte': date_start, '$lte': date_end}}
            ]
        }))
        
        for order in orders:
            activities.append(format_order_for_date(order, date_start, date_end))
    
    elif filter_type == 'modified':
        orders = list(orders_collection.find({
            'edit_tag': {'$exists': True},
            '$or': [
                {'date': {'$gte': date_start, '$lte': date_end}},
                {'edit_tag.modification_date': {'$gte': date_start, '$lte': date_end}}
            ]
        }))
        
        for order in orders:
            activities.append(format_order_for_date(order, date_start, date_end))
    
    elif filter_type == 'previous':
        orders = list(orders_collection.find({
            'date': {'$lt': date_start},
            'payment_history.date': {'$gte': date_start, '$lte': date_end}
        }))
        
        for order in orders:
            activities.append(format_order_for_date(order, date_start, date_end))
    
    # Sort by timestamp (newest first)
    activities.sort(key=lambda x: x.get('timestamp', ''), reverse=True)
    
    # Paginate
    total = len(activities)
    total_pages = (total + per_page - 1) // per_page if total > 0 else 1
    start_idx = (page - 1) * per_page
    end_idx = start_idx + per_page
    paginated = activities[start_idx:end_idx]
    
    return {
        'data': paginated,
        'pagination': {
            'current_page': page,
            'per_page': per_page,
            'total': total,
            'total_pages': total_pages,
            'has_prev': page > 1,
            'has_next': page < total_pages
        }
    }

def calculate_stats(activities, filter_type):
    """Calculate stats - orders only, no expenses in order count"""
    stats = {
        'total_orders': 0,
        'gross_total': 0,
        'debt': 0,
        'expenses': 0,
        'modified_orders': 0,
        'net': 0
    }
    
    order_ids_seen = set()
    
    for activity in activities:
        if activity['type'] == 'order':
            # Count unique orders only
            if activity['receipt_id'] not in order_ids_seen:
                order_ids_seen.add(activity['receipt_id'])
                stats['total_orders'] += 1
                
                # Gross total = sum of all order totals (what should be collected)
                stats['gross_total'] += activity.get('total', 0)
                
                # Debt = unpaid balance
                if activity.get('balance', 0) > 0:
                    stats['debt'] += activity['balance']
                
                # Modified count
                if activity.get('is_modified'):
                    stats['modified_orders'] += 1
        
        elif activity['type'] == 'expense':
            stats['expenses'] += activity['amount']
    
    # NET = Money collected - Expenses = (Gross - Debt) - Expenses
    stats['net'] = (stats['gross_total'] - stats['debt']) - stats['expenses']
    
    return stats


def calculate_filter_counts(orders_collection, expenses_collection):
    """Calculate counts for filter badges"""
    return {
        'all': orders_collection.count_documents({}) + expenses_collection.count_documents({}),
        'pending': orders_collection.count_documents({'balance': {'$gt': 0}}),
        'completed': orders_collection.count_documents({'balance': {'$lte': 0}}),
        'expenses': expenses_collection.count_documents({}),
        'gateway': orders_collection.count_documents({
            '$or': [
                {'payment_type': 'mpesa'},
                {'payment_breakdown.mpesa': {'$exists': True}},
                {'payment_history.payment_type': 'mpesa'}
            ]
        }),
        'modified': orders_collection.count_documents({'edit_tag': {'$exists': True}}),
        'previous': orders_collection.count_documents({
            'date': {'$lt': datetime.now(NAIROBI_TZ).replace(hour=0, minute=0, second=0)},
            'payment_history.date': {'$gte': datetime.now(NAIROBI_TZ).replace(hour=0, minute=0, second=0)}
        })
    }


# orders definitions
def create_notification(order_data, notification_type='order_created'):
    """Create a notification from order data"""
    try:
        # Calculate total amount from items
        items = order_data.get('items', [])
        total_amount = 0
        
        i = 0
        while i < len(items):
            if items[i] == 'product':
                quantity = items[i + 3] if i + 3 < len(items) else 0
                price = items[i + 5] if i + 5 < len(items) else 0
                total_amount += quantity * price
                i += 6
            else:
                i += 1
        
        # Determine status based on payment
        payment = order_data.get('payment', 0)
        balance = order_data.get('balance', 0)
        
        if balance == 0:
            status = 'paid'
        elif payment > 0 and balance > 0:
            status = 'partial'
        else:
            status = 'credit'
        
        # Get next notification ID
        last_notif = notifications_collection.find_one(sort=[('notification_id', -1)])
        if last_notif and last_notif.get('notification_id'):
            last_id = int(last_notif['notification_id'].replace('NOTIF', ''))
            next_id = f"NOTIF{str(last_id + 1).zfill(6)}"
        else:
            next_id = 'NOTIF000001'
        
        # Create notification document
        notification = {
            'notification_id': next_id,
            'type': notification_type,
            'category': 'orders',  # Added category for grouping
            'receipt_id': order_data.get('receipt_id'),
            'salesperson_name': order_data.get('salesperson_name', 'N/A'),
            'shop_name': order_data.get('shop_name', 'Unknown Shop'),
            'amount': total_amount,
            'status': status,
            'order_type': order_data.get('order_type', 'wholesale'),
            'created_at': datetime.now(NAIROBI_TZ),
            'read': False
            # NO 'cleared' field - we keep all notifications permanently
        }
        
        # Insert notification
        notifications_collection.insert_one(notification)
        logger.info(f"Notification {next_id} created for order {order_data.get('receipt_id')}")
        
        return next_id
        
    except Exception as e:
        logger.error(f"Error creating notification: {e}")
        return None


# Payment notification helper - UPDATED
def create_payment_notification(receipt_id, payment_amount, payment_type):
    """Create notification when payment is received"""
    try:
        order = orders_collection.find_one({'receipt_id': receipt_id})
        if not order:
            return None
        
        # Get next notification ID
        last_notif = notifications_collection.find_one(sort=[('notification_id', -1)])
        if last_notif and last_notif.get('notification_id'):
            last_id = int(last_notif['notification_id'].replace('NOTIF', ''))
            next_id = f"NOTIF{str(last_id + 1).zfill(6)}"
        else:
            next_id = 'NOTIF000001'
        
        # Determine new status
        new_balance = order.get('balance', 0) - payment_amount
        if new_balance <= 0:
            status = 'paid'
        else:
            status = 'partial'
        
        notification = {
            'notification_id': next_id,
            'type': 'payment_received',
            'category': 'orders',  # Added category
            'receipt_id': receipt_id,
            'salesperson_name': order.get('salesperson_name', 'N/A'),
            'shop_name': order.get('shop_name', 'Unknown Shop'),
            'amount': payment_amount,
            'status': status,
            'order_type': order.get('order_type', 'wholesale'),
            'created_at': datetime.now(NAIROBI_TZ),
            'read': False
        }
        
        notifications_collection.insert_one(notification)
        logger.info(f"Payment notification {next_id} created for order {receipt_id}")
        
        return next_id
        
    except Exception as e:
        logger.error(f"Errors creating payment notification: {e}")
        return None

    
def clean_doc(doc):
    """Convert ObjectId and other non-serializable types to strings."""
    cleaned = {}
    for k, v in doc.items():
        if isinstance(v, ObjectId):
            cleaned[k] = str(v)
        elif isinstance(v, dict):
            cleaned[k] = clean_doc(v)  # recursive
        elif isinstance(v, list):
            cleaned[k] = [clean_doc(i) if isinstance(i, dict) else str(i) if isinstance(i, ObjectId) else i for i in v]
        else:
            cleaned[k] = v
    return cleaned

def parse_mongo_date(date_str):
    """Handle MongoDB string dates like '2025-10-22T13:50:56.123Z' or '2025-10-22 13:50:56'"""
    if not date_str:
        return None
    try:
        if isinstance(date_str, str):
            # Try ISO format first
            return parser.parse(date_str).astimezone(NAIROBI_TZ)
        elif isinstance(date_str, datetime):
            return date_str.astimezone(NAIROBI_TZ)
    except:
        try:
            # Fallback: split and reconstruct
            return parser.parse(str(date_str).split('.')[0]).astimezone(NAIROBI_TZ)
        except:
            pass
    return None

def serialize_mongo_doc(doc):
    """Convert MongoDB document to JSON-serializable dict"""
    if not doc:
        return doc
    
    doc = doc.copy()
    if '_id' in doc:
        doc['_id'] = str(doc['_id'])
    if 'date' in doc:
        doc['date'] = process_date(doc.get('date'))
    return doc




def calculate_dashboard_stats(orders, retail_collection, today_start, today_end):
    """
    Calculate dashboard statistics for sales, debts, and order counts.
    🆕 ENHANCED: Now properly handles daily expenses and debts
    """
    retail_sales_today = 0.0
    wholesale_sales_today = 0.0
    total_debts = 0.0
    open_orders_count = 0
    closed_orders_count = 0
    retail_open_orders = 0
    retail_closed_orders = 0
    wholesale_open_orders = 0
    wholesale_closed_orders = 0
    
    # 🆕 NEW: Calculate TODAY's expenses only (resets daily)
    total_expenses = sum(
        float(doc.get('amount', 0)) 
        for doc in db["expenses"].find({
            'date': {'$gte': today_start, '$lt': today_end}
        })
    )
    
    for order in orders:
        order_dict = order
        order_date = process_date(order_dict.get('date'))
        closed_date = process_date(order_dict.get('closed_date'))
        order_type = order_dict.get('order_type', 'wholesale')
        payment = float(order_dict.get('payment', 0))
        balance = float(order_dict.get('balance', 0))
        payment_history = order_dict.get('payment_history', [])
        status = order_dict.get('status', 'pending' if balance > 0 else 'completed')
        
        # Prevent closed_date for pending orders
        if balance > 0:
            closed_date = None
        
        # Update open/closed counts for today
        if order_date and order_date >= today_start and order_date < today_end:
            if status == 'pending' or balance > 0:
                open_orders_count += 1
                if order_type in ['retail', 'app']:
                    retail_open_orders += 1
                else:
                    wholesale_open_orders += 1
            else:
                closed_orders_count += 1
                if order_type in ['retail', 'app']:
                    retail_closed_orders += 1
                else:
                    wholesale_closed_orders += 1
        
        # 🆕 MODIFIED: Only count debts created TODAY (resets daily)
        if balance > 0 and order_date and order_date >= today_start and order_date < today_end:
            total_debts += balance
        
        # Calculate today's sales from payment_history
        for payment_entry in payment_history:
            payment_date = process_date(payment_entry.get('date'))
            payment_amount = float(payment_entry.get('amount', 0))
            
            if payment_date and payment_date >= today_start and payment_date < today_end and payment_amount > 0:
                if order_type in ['retail', 'app']:
                    retail_sales_today += payment_amount
                    logger.debug(f"Order {order['_id']} (payment today, type {order_type}): Added {payment_amount} to retail_sales_today")
                else:
                    wholesale_sales_today += payment_amount
                    logger.debug(f"Order {order['_id']} (payment today, type {order_type}): Added {payment_amount} to wholesale_sales_today")
    
    # Add retail collection amounts
    retail_collection_total = sum(
        float(r.get('amount', 0))
        for r in retail_collection
        if float(r.get('amount', 0)) > 0
    )
    
    if retail_collection_total > 0:
        retail_sales_today += retail_collection_total
        logger.debug(f"Added {retail_collection_total} from retail collection to retail_sales_today")
    
    # 🆕 MODIFIED: Subtract expenses from total sales
    total_sales_today = retail_sales_today + wholesale_sales_today - total_expenses
    
    return {
        'total_sales_today': round(total_sales_today, 2),
        'retail_sales_today': round(retail_sales_today, 2),
        'wholesale_sales_today': round(wholesale_sales_today, 2),
        'total_debt': round(total_debts, 2),  # 🆕 Changed key name for consistency
        'total_expenses': round(total_expenses, 2),  # 🆕 Added to return
        'open_orders_count': open_orders_count,
        'closed_orders_count': closed_orders_count,
        'retail_open_orders': retail_open_orders,
        'retail_closed_orders': retail_closed_orders,
        'wholesale_open_orders': wholesale_open_orders,
        'wholesale_closed_orders': wholesale_closed_orders
    }


# stock functions
def update_stock_version():
    """Increment the stock version in MongoDB with detailed logging."""
    client_log_messages = []
    logger.info("[UPDATE_STOCK_VERSION] Starting stock version update")
    client_log_messages.append("Starting stock version update")

    try:
        # Fetch and update version
        version_doc = metadata_collection.find_one_and_update(
            {'_id': 'stock_version'},
            {'$inc': {'version': 1}},
            upsert=True,
            return_document=True
        )
        logger.info("[UPDATE_STOCK_VERSION] Fetched stock_version document")
        client_log_messages.append("Fetched stock_version document")

        # Determine current version
        current_version = version_doc.get('version', 1) if version_doc else 1
        logger.info(f"[UPDATE_STOCK_VERSION] Current version: {current_version}, type: {type(current_version)}")
        client_log_messages.append(f"Current version: {current_version}, type: {type(current_version).__name__}")

        # Validate version type
        if not isinstance(current_version, (int, float)):
            logger.warning(f"[UPDATE_STOCK_VERSION] Invalid version type: {type(current_version)}, resetting to 1")
            client_log_messages.append(f"Invalid version type: {type(current_version).__name__}, resetting to 1")
            current_version = 1
            metadata_collection.update_one(
                {'_id': 'stock_version'},
                {'$set': {'version': current_version}},
                upsert=True
            )

        logger.info(f"[UPDATE_STOCK_VERSION] Successfully updated stock_version to {current_version}")
        client_log_messages.append(f"Successfully updated stock_version to {current_version}")
        return current_version, client_log_messages

    except Exception as e:
        error_message = f"Error in update_stock_version: {str(e)}\n{traceback.format_exc()}"
        logger.error("[UPDATE_STOCK_VERSION] %s", error_message)
        client_log_messages.append(error_message)
        raise

def clear_stock_cache_logic():
    """Clear stock cache logic without HTTP overhead."""
    client_log_messages = []
    logger.info("[CLEAR_STOCK_CACHE] Clearing cache: user=%s", session.get('user', {}).get('email', 'unknown'))
    client_log_messages.append(f"Clearing cache: user={session.get('user', {}).get('email', 'unknown')}")

    try:
        logger.info("[CLEAR_STOCK_CACHE] Calling update_stock_version")
        client_log_messages.append("Calling update_stock_version")
        new_version, update_logs = update_stock_version()
        client_log_messages.extend(update_logs)

        logger.info("[CLEAR_STOCK_CACHE] Clearing stock_cache")
        client_log_messages.append("Clearing stock_cache")
        stock_cache['data'] = None
        stock_cache['version'] = None
        stock_cache['timestamp'] = None
        logger.info("[CLEAR_STOCK_CACHE] Stock cache cleared successfully")
        client_log_messages.append("Stock cache cleared successfully")
        return True, client_log_messages
    except Exception as e:
        error_message = f"Error in clear_stock_cache_logic: {str(e)}\n{traceback.format_exc()}"
        logger.error("[CLEAR_STOCK_CACHE] %s", error_message)
        client_log_messages.append(error_message)
        return False, client_log_messages

   #log functions 
def log_stock_change(product_type, subtype, change_type, quantity, price_per_unit):
    """Log stock changes. Collection: inventory_logs"""
    now = datetime.now(NAIROBI_TZ)
    db["inventory_logs"].insert_one({
        'category': product_type,     # Renamed for dashboard consistency
        'item': subtype,             # Renamed for dashboard consistency
        'action': change_type,
        'quantity': quantity,
        'price': price_per_unit,
        'date': now,                 # Use actual datetime object, not string
        'user_id': session.get('user', {}).get('uid', 'System'),
        # Adding a description helps the AJAX feed show a clean text line
        'description': f"{change_type.replace('_', ' ').title()}: {subtype} ({quantity})"
    })

def log_user_action(action_type, details):
    """Log user actions. Collection: audit_trail"""
    user_name = f"{session['user']['first_name']} {session['user']['last_name']}" if 'user' in session else "Unknown User"
    db["audit_trail"].insert_one({
        'user_name': user_name,
        'action': action_type,       # Renamed for consistency
        'details': details,
        'date': datetime.now(NAIROBI_TZ) # Use datetime object
    })

# receipts function
def get_next_receipt_id():
    counter = metadata_collection.find_one_and_update(
        {'_id': 'receipt_counter'},
        {'$inc': {'last_id': 1}},
        upsert=True,
        return_document=True
    )
    return f"DRE{counter.get('last_id', 1):06d}52"  
# orders function

def process_items(items: Union[List, str]) -> int:
    """Calculate the number of items or total quantity from a list or string."""
    if not items:
        return 0
    try:
        if isinstance(items, str):
            items = json.loads(items)
        if not isinstance(items, list):
            return 0
        
        if items and isinstance(items[0], dict):  # App order format: [{'name': ..., 'quantity': ...}, ...]
            return sum(item.get('quantity', 0) for item in items if isinstance(item.get('quantity'), (int, float)))
        
        # Web order format: ['product', name, 'quantity', qty, 'price', price, ...]
        count = 0
        i = 0
        while i < len(items):
            if items[i] == 'product':
                count += 1
                i += 6  # Skip product tuple
            else:
                i += 1
        return count
    except (json.JSONDecodeError, IndexError, TypeError):
        return 0

def get_safe_name(user_val):
    if isinstance(user_val, dict):
        fname = user_val.get('first_name', '')
        lname = user_val.get('last_name', '')
        return f"{fname} {lname}".strip() or user_val.get('email', 'Unknown User')
    return user_val if isinstance(user_val, str) else 'Unknown User'

def process_order(doc):
    # doc is ALREADY a dict from MongoDB - NO .to_dict() needed!
    order_dict = doc  # Just use doc directly
    
    balance = float(order_dict.get('balance', 0))
    closed_date = process_date(order_dict.get('closed_date')) if order_dict.get('closed_date') else None
    status = order_dict.get('status', 'pending' if balance > 0 else 'completed')
    if balance > 0:
        closed_date = None
    salesperson_name = resolve_salesperson_name(order_dict)

    # Safely process payment_history - skip anything that isn't a dict
    payment_history = []
    for ph in order_dict.get('payment_history', []):
        if isinstance(ph, dict):
            payment_history.append({
                'amount': float(ph.get('amount', 0)),
                'date': process_date(ph.get('date')),
                'payment_type': ph.get('payment_type', '')
            })
        # if ph is a string or anything else, just skip it
    
    return {
        'receipt_id': order_dict.get('receipt_id', str(doc.get('_id', ''))),
        'salesperson_name': salesperson_name,
        'salesperson_name_lower': salesperson_name.lower(),
        'client_name': order_dict.get('shop_name', 'Unknown Client'),
        'shop_name': order_dict.get('shop_name', 'Unknown Shop'),
        'shop_name_lower': order_dict.get('shop_name_lower', 'unknown shop'),
        'items': json.dumps(order_dict.get('items', [])),
        'photoUrl': order_dict.get('photoUrl', ''),
        'payment': float(order_dict.get('payment', 0)),
        'balance': balance,
        'date': process_date(order_dict.get('date')),
        'closed_date': closed_date,
        'order_type': order_dict.get('order_type', 'wholesale'),
        'payment_type': order_dict.get('payment_type', ''),
        'payment_history': payment_history,
        'edit_tag': order_dict.get('edit_tag'),
        'notes': order_dict.get('notes', ''),
        'status': status,
        'user_id': order_dict.get('user_id', '')
    }

def format_currency(value):
    try:
        return f"{float(value):,.2f}"  # Format as currency with 2 decimal places
    except (ValueError, TypeError):
        return value  # Return as-is if conversion fails

app.jinja_env.filters['format_currency'] = format_currency 

def format_number(value):
    try:
        return f"{int(float(value)):,}"  # Format as integer with commas (no decimals)
    except (ValueError, TypeError):
        return str(value)  # Return as-is if conversion fails
app.jinja_env.filters['format_number'] = format_number

#clients 
def update_clients_counter(change, context):
    count = sum(1 for _ in db.collection('clients').stream())
    db.collection('metadata').document('clients_counter').set({'count': count})

# resolving persons name function
def resolve_salesperson_name(order_dict):
    """Resolve salesperson_name from users collection if Anonymous for app orders."""
    salesperson_name = order_dict.get('salesperson_name', 'N/A')
    if salesperson_name.lower() == 'anonymous' and order_dict.get('order_type') == 'app':
        user_id = order_dict.get('user_id')
        if user_id:
            try:
                user_doc = db.collection('users').document(user_id).get()
                if user_doc.exists:
                    user_data = user_doc.to_dict()
                    name = f"{user_data.get('first_name', '')} {user_data.get('last_name', '')}".strip()
                    if name:
                        salesperson_name = name
                        logger.debug(f"Resolved Anonymous to {name} for user_id {user_id}")
                    else:
                        logger.warning(f"User {user_id} has no first_name or last_name")
                else:
                    logger.warning(f"No user found for user_id {user_id}")
            except Exception as e:
                logger.error(f"Error resolving salesperson_name for user_id {user_id}: {e}")
    return salesperson_name

# date and time functions
def expire_date_days_left(date_str):
    """Calculate days until expiry date, handling all edge cases."""
    if not date_str or date_str in [None, "", "0000-00-00 00:00:00"]:
        return None
    try:
        expiry_date = datetime.strptime(date_str, "%Y-%m-%d")
        today = datetime.now(NAIROBI_TZ).replace(hour=0, minute=0, second=0, microsecond=0)
        days_left = (expiry_date - today).days
        return max(days_left, 0)  # Ensure no negative days
    except (ValueError, TypeError):
        return None
 
def process_date(date_value):
    """Convert a date value to a datetime object in Nairobi timezone."""
    if date_value is None or date_value == '':
        return None
    try:
        if date_value.tzinfo is None:
            return UTC.localize(date_value).astimezone(NAIROBI_TZ)  # Correct

        if isinstance(date_value, datetime):
            if date_value.tzinfo is None:
                return NAIROBI_TZ.localize(date_value)
            return date_value.astimezone(NAIROBI_TZ)

        elif isinstance(date_value, str):
            # Try isoformat first (what .isoformat() and MongoDB actually produce)
            # e.g. '2026-01-20T06:27:12.991+00:00' or '2026-01-20T06:27:12.991Z'
            try:
                # Replace Z with +00:00 so fromisoformat can handle it
                cleaned = date_value.replace('Z', '+00:00')
                parsed_date = datetime.fromisoformat(cleaned)
                return parsed_date.astimezone(NAIROBI_TZ)
            except ValueError:
                pass

            # Fallback: plain date string '2026-01-02'
            try:
                parsed_date = datetime.strptime(date_value, '%Y-%m-%d')
                return NAIROBI_TZ.localize(parsed_date)
            except ValueError:
                pass

            # Nothing worked
            logger.error(f"Could not parse date string: {date_value}")
            return None

        else:
            logger.error(f"Unexpected date type: {type(date_value)} - value: {date_value}")
            return None

    except Exception as e:
        logger.error(f"Error processing date: {str(e)} - value: {date_value}")
        return None

def parse_items(items_list):
    """Parse items list into clean product string"""
    products = []
    i = 0
    while i < len(items_list):
        if items_list[i] == 'product' and i + 5 < len(items_list):
            name = items_list[i+1]
            qty = items_list[i+3]
            products.append(f"{name} x{qty}")
            i += 6
        else:
            i += 1
    return ", ".join(products) if products else "N/A"

def generate_pdf_export(orders, expenses, time_filter, status_filter):
    buffer = BytesIO()
    # Use a standard font throughout for that "Balance Sheet" look
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    
    # 1. PROFESSIONAL HEADER
    p.setFont("Helvetica-Bold", 12)
    p.drawString(50, height - 40, "DREAMLAND ZERO - CONSOLIDATED FINANCIAL REPORT")
    
    # Calculate Date Range for Header
    now = datetime.now(NAIROBI_TZ)
    start_date = "Beginning"
    if time_filter == 'day': start_date = now.strftime('%d/%m/%Y 00:00')
    elif time_filter == 'week': start_date = (now - timedelta(days=now.weekday())).strftime('%d/%m/%Y')
    elif time_filter == 'month': start_date = now.strftime('01/%m/%Y')
    
    p.setFont("Helvetica", 8)
    date_range_str = f"PERIOD: {start_date} TO {now.strftime('%d/%m/%Y %H:%M')}"
    p.drawString(50, height - 52, date_range_str)
    p.line(50, height - 60, width - 50, height - 60)

    # 2. SALES TABLE (High Density / Excel Style)
    y = height - 80
    col_heads = [
        (50, "Receipt"), (105, "Date"), (150, "Salesperson"), 
        (220, "Items"), (360, "Value"), (420, "Paid"), (480, "Debt")
    ]
    
    p.setFont("Helvetica-Bold", 8)
    for x, text in col_heads:
        p.drawString(x, y, text)
    
    y -= 4
    p.line(50, y, width - 50, y)
    y -= 10

    total_sales_val = 0
    total_collected = 0
    total_debt = 0

    p.setFont("Helvetica", 7) # Small, uniform font for data
    for order in orders:
        if y < 100: # Standard break for summary space
            p.showPage()
            y = height - 50
            p.setFont("Helvetica", 7)

        # Logic for Value & History
        items = order.get('items', [])
        order_total = 0
        idx = 0
        while idx < len(items):
            if items[idx] == 'product' and idx + 5 < len(items):
                order_total += (items[idx+3] * items[idx+5])
                idx += 6
            else: idx += 1

        paid_to_date = sum(item.get('amount', 0) for item in order.get('payment_history', []))
        debt = max(0, order_total - paid_to_date)

        # Draw Row (Minified Spacing)
        p.drawString(50, y, str(order.get('receipt_id', 'N/A')))
        p.drawString(105, y, order.get('date').strftime('%d/%m/%y') if order.get('date') else "-")
        p.drawString(150, y, str(order.get('salesperson_name', 'User'))[:15])
        
        prod_text = parse_items(items)
        p.drawString(220, y, prod_text[:35] + ('..' if len(prod_text) > 35 else ''))
        
        # Align numbers to the right visually by padding or fixed x
        p.drawRightString(405, y, f"{order_total:,.0f}")
        p.drawRightString(465, y, f"{paid_to_date:,.0f}")
        p.drawRightString(525, y, f"{debt:,.0f}")
        
        total_sales_val += order_total
        total_collected += paid_to_date
        total_debt += debt
        y -= 9 # Tight Excel-like row height

    # 3. EXPENSES TABLE (Same Header/Font Size)
    y -= 15
    p.setFont("Helvetica-Bold", 8)
    p.drawString(50, y, "EXPENSES")
    y -= 4
    p.line(50, y, 530, y)
    y -= 10
    
    total_expenses = 0
    p.setFont("Helvetica", 7)
    for exp in expenses:
        if y < 60:
            p.showPage()
            y = height - 50
        
        amt = exp.get('amount', 0)
        p.drawString(50, y, f"[{exp.get('category', 'Other')}]")
        p.drawString(150, y, exp.get('description', '')[:65])
        p.drawRightString(525, y, f"{amt:,.0f}")
        total_expenses += amt
        y -= 9

    # 4. PROFESSIONAL BALANCE SUMMARY
    y -= 25
    if y < 80: p.showPage(); y = height - 50
    
    # Box for Summary
    p.setStrokeColorRGB(0.8, 0.8, 0.8)
    p.rect(50, y - 45, 480, 60, fill=0)
    p.setStrokeColorRGB(0, 0, 0)
    
    y -= 5
    p.setFont("Helvetica-Bold", 9)
    p.drawString(60, y, "FINANCIAL SUMMARY")
    
    p.setFont("Helvetica", 8)
    y -= 15
    p.drawString(60, y, f"Total Sales Value: {total_sales_val:,.2f}")
    p.drawString(230, y, f"Cash Collected: {total_collected:,.2f}")
    p.drawString(400, y, f"Expenses: ({total_expenses:,.2f})")
    
    y -= 15
    p.setFont("Helvetica-Bold", 9)
    p.drawString(60, y, f"ACCOUNTS RECEIVABLE (DEBT): {total_debt:,.2f}")
    
    # Final Net Profit/Cash Position
    net_cash = total_collected - total_expenses
    p.drawRightString(520, y, f"NET CASH POSITION: KES {net_cash:,.2f}")
    
    p.save()
    buffer.seek(0)
    
    filename = f"Report_{now.strftime('%Y%m%d_%H%M')}.pdf"
    return Response(
        buffer.getvalue(),
        mimetype='application/pdf',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )

def validate_items(items):
    """Validate item structure and values"""
    if not items or len(items) == 0:
        return False, "At least one item required"
    
    for item in items:
        if item.get('quantity', 0) <= 0:
            return False, f"Invalid quantity for {item.get('product', 'item')}"
        if item.get('price', 0) <= 0:
            return False, f"Invalid price for {item.get('product', 'item')}"
    
    return True, None

def calculate_order_total(items):
    """Calculate total from items array"""
    return sum(item['price'] * item['quantity'] for item in items)

def generate_docx_export(orders, time_filter, status_filter):
    """Generate DOCX export - Excel style table"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return jsonify({'error': 'python-docx not installed'}), 500
    
    doc = Document()
    
    # Header
    title = doc.add_heading('DREAMLAND SALES REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'{time_filter.upper()} | {status_filter.upper()} | {datetime.now(NAIROBI_TZ).strftime("%d/%m/%Y %H:%M")}')
    
    # Table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid'
    
    # Headers
    headers = table.rows[0].cells
    headers[0].text = 'Receipt'
    headers[1].text = 'Date'
    headers[2].text = 'Salesperson'
    headers[3].text = 'Products'
    headers[4].text = 'Total'
    headers[5].text = 'Paid'
    headers[6].text = 'Debt'
    
    for cell in headers:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    
    # Data
    total_sales = 0
    total_paid = 0
    total_debt = 0
    
    for order in orders:
        row = table.add_row().cells
        
        total = order.get('payment', 0) + order.get('balance', 0)
        paid = order.get('payment', 0)
        debt = order.get('balance', 0)
        
        row[0].text = str(order.get('receipt_id', 'N/A'))
        row[1].text = order.get('date', datetime.now()).strftime('%d/%m/%Y')
        row[2].text = str(order.get('salesperson_name', 'N/A'))
        row[3].text = parse_items(order.get('items', []))
        row[4].text = f"{total:,.0f}"
        row[5].text = f"{paid:,.0f}"
        row[6].text = f"{debt:,.0f}" if debt > 0 else "-"
        
        # Set font size for data rows
        for cell in row:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
        
        total_sales += total
        total_paid += paid
        total_debt += debt
    
    # Summary row
    summary = table.add_row().cells
    summary[0].text = f'TOTAL ({len(orders)} orders)'
    summary[3].text = ''
    summary[4].text = f"{total_sales:,.0f}"
    summary[5].text = f"{total_paid:,.0f}"
    summary[6].text = f"{total_debt:,.0f}"
    
    for cell in summary:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    
    # Save
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return Response(
        buffer.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': f'attachment; filename=sales_{datetime.now(NAIROBI_TZ).strftime("%Y%m%d_%H%M%S")}.docx'}
    )

# ============================================================================
# ROUTES
# ============================================================================

@app.before_request
def detect_abnormal_session_end():
    if request.endpoint in ['static', 'auth', 'splash']:
        return

    current_user = get_current_user()
    if not current_user:
        return

    user_email = current_user.get('email')
    if not user_email:
        return

    current_login_time_str = session.get('login_time')
    if not current_login_time_str:
        return

    current_login_time = datetime.fromisoformat(current_login_time_str)
    if current_login_time.tzinfo is None:
        current_login_time = NAIROBI_TZ.localize(current_login_time)

    last_login_log = db['session_logs'].find_one(
        {"email": user_email, "action": "login_success"},
        sort=[("timestamp", -1)]
    )
    if not last_login_log:
        return

    last_db_login_time = last_login_log.get('timestamp')
    if not last_db_login_time:
        return

    if last_db_login_time.tzinfo is None:
        last_db_login_time = NAIROBI_TZ.localize(last_db_login_time)

    if abs((last_db_login_time - current_login_time).total_seconds()) < 2:
        return

    existing_logout = db['session_logs'].find_one({
        "email": user_email,
        "action": "logout",
        "timestamp": {"$gt": last_db_login_time}
    })
    if existing_logout:
        return

    db['session_logs'].insert_one({
        'email':       user_email,
        'action':      'logout',
        'logout_type': 'abnormal',
        'reason':      'session_expired_or_browser_closed',
        'last_login':  last_db_login_time,
        'timestamp':   datetime.now(NAIROBI_TZ),
        'ip_address':  request.remote_addr,
        'user_agent':  request.headers.get('User-Agent')
    })


@app.route('/api/health', methods=['GET'])
def health_check():
    """Check application and database health"""
    try:
        db_manager.client.admin.command('ping')
        return jsonify({
            'status': 'healthy',
            'database': {
                'connected': True,
                'host': Config.MY_HOST,
                'peer_online': db_manager.peer_online,
                'mode': 'STANDALONE'
            },
            'timestamp': datetime.now(NAIROBI_TZ).isoformat()
        }), 200
    except Exception as e:
        return jsonify({'status': 'unhealthy', 'error': str(e)}), 503
    
@app.route('/')
def splash():
    if 'user' in session:
        return redirect(url_for('dashboard'))
    return render_template('splash.html')


def login_required(f):
    """Enhanced login_required decorator"""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session:
            flash('Please log in to access this page.', 'warning')
            return redirect(url_for('auth', next=request.url))
        return f(*args, **kwargs)
    return decorated_function

# ============================================================================
# AUTHENTICATION ROUTE (Enhanced)
# ============================================================================

@app.route('/auth', methods=['GET', 'POST'])
def auth():
    """
    Unified authentication endpoint for login and signup
    - Handles form validation
    - Tracks login attempts
    - **NEW: Tracks last_login timestamps**
    """
    if 'user' in session:
        return redirect(url_for('dashboard'))
    
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        # ========================================
        # SIGNUP FLOW
        # ========================================
        if form_type == 'signup':
            email = request.form.get('email')
            first_name = request.form.get('first_name')
            last_name = request.form.get('last_name')
            phone = request.form.get('phone')
            password = request.form.get('password')
            confirm_password = request.form.get('confirmPassword')
            role = request.form.get('role', 'user')

            print(f"DEBUG first_name: {repr(first_name)}")
            print(f"DEBUG last_name: {repr(last_name)}")

            # Server-side validation
            if not first_name or not last_name:
                return jsonify({"status": "error", "error": "Name fields are required"}), 400
            
            if not re.match(r"^[A-Za-z]+$", first_name) or not re.match(r"^[A-Za-z]+$", last_name):
                return jsonify({"status": "error", "error": "Names must contain only letters"}), 400
            
            if password != confirm_password:
                return jsonify({"status": "error", "error": "Passwords don't match."}), 400
            
            if not re.match(r"[^@]+@[^@]+\.[^@]+", email):
                return jsonify({"status": "error", "error": "Invalid email format."}), 400
            
            if not re.match(r"[0-9]{10,15}", phone):
                return jsonify({"status": "error", "error": "Phone number should be 10-15 digits."}), 400
            
            if not re.match(r"(?=.*\d)(?=.*[a-z])(?=.*[A-Z]).{8,}", password):
                return jsonify({"status": "error", "error": "Password must be at least 8 characters, including a number, an uppercase letter, and a lowercase letter."}), 400
            
            try:
                # Check if email exists
                if users_collection.find_one({"email": email}):
                    return jsonify({"status": "error", "error": "Email already exists. Try logging in."}), 400
                
                # Hash password and create user
                hashed_password = generate_password_hash(password, method='pbkdf2:sha256')
                users_collection.insert_one({
                    "email": email,
                    "first_name": first_name,
                    "last_name": last_name,
                    "phone": phone,
                    "password": hashed_password,
                    "role": role,
                    "status": "pending",
                    "created_at": datetime.now(NAIROBI_TZ),
                    "last_login": None,  # 🆕 NEW: Initialize last_login
                    "login_history": []  # 🆕 NEW: Track login history
                })
                
                # Log signup
                db['session_logs'].insert_one({
                    'email': email,
                    'action': 'signup',
                    'timestamp': datetime.now(NAIROBI_TZ),
                    'ip_address': request.remote_addr,
                    'user_agent': request.headers.get('User-Agent')
                })
                
                return jsonify({
                    "status": "success",
                    "message": "Signup successful! Please log in."
                }), 200
                
            except Exception as e:
                return jsonify({"status": "error", "error": f"Signup failed: {str(e)}"}), 500
        
        # ========================================
        # LOGIN FLOW (🆕 ENHANCED)
        # ========================================
        elif form_type == 'login':
            email = request.form.get('email')
            password = request.form.get('password')
            
            try:
                user = users_collection.find_one({"email": email})
                
                # Check if user exists
                if not user:
                    # Log failed attempt
                    db['session_logs'].insert_one({
                        'email': email,
                        'action': 'login_failed',
                        'reason': 'user_not_found',
                        'timestamp': datetime.now(NAIROBI_TZ),
                        'ip_address': request.remote_addr,
                        'user_agent': request.headers.get('User-Agent')
                    })
                    return jsonify({"status": "error", "error": "No account with that email."}), 401
                
                # Check password
                if not check_password_hash(user["password"], password):
                    # Log failed attempt
                    db['session_logs'].insert_one({
                        'email': email,
                        'action': 'login_failed',
                        'reason': 'wrong_password',
                        'timestamp': datetime.now(NAIROBI_TZ),
                        'ip_address': request.remote_addr,
                        'user_agent': request.headers.get('User-Agent')
                    })
                    return jsonify({"status": "error", "error": "Wrong password. Try again."}), 401
                # Check if blocked
                if user.get('status') == 'blocked':
                    db['session_logs'].insert_one({
                        'email': email,
                        'action': 'login_failed',
                        'reason': 'account_blocked',
                        'timestamp': datetime.now(NAIROBI_TZ),
                        'ip_address': request.remote_addr,
                        'user_agent': request.headers.get('User-Agent')
                    })
                    return jsonify({"status": "error", "error": "Your account has been suspended. Contact support."}), 403

                # Check if pending approval
                if user.get('status') == 'pending':
                    return jsonify({"status": "pending", "redirect": f"/awaiting?email={email}"}), 200
                
                # 🆕 NEW: Capture previous last_login before updating
                previous_last_login = user.get('last_login')
                current_login_time = datetime.now(NAIROBI_TZ)
                
                # 🆕 NEW: Update user's last_login in database
              
                # Initialize session
                session['user'] = {
                    "first_name": user.get("first_name"),
                    "last_name": user.get("last_name"),
                    "email": user.get("email"),
                    "role": user.get("role", "user"),
                    "user_id": str(user["_id"])
                }
                session['last_activity'] = current_login_time.isoformat()
                session['login_time'] = current_login_time.isoformat()
                session['previous_last_login'] = previous_last_login.isoformat() if previous_last_login else None  # 🆕 NEW
                session.permanent = True
                
                # Log successful login
                db['session_logs'].insert_one({
                    'email': email,
                    'action': 'login_success',
                    'timestamp': current_login_time,
                    'ip_address': request.remote_addr,
                    'user_agent': request.headers.get('User-Agent'),
                    'previous_login': previous_last_login  # 🆕 NEW: Track previous login
                })
                
                # Redirect to intended page or dashboard
                next_page = request.args.get('next', url_for('dashboard'))
                return jsonify({
                    "status": "success",
                    "message": "Login successful!",
                    "redirect": next_page
                }), 200
                
            except Exception as e:
                return jsonify({"status": "error", "error": f"Login failed: {str(e)}"}), 500
    
    return render_template('auth.html', signup_success=request.args.get('signup_success', False))



# ============================================================================
# ADMIN: View Session Logs
# ============================================================================

def get_current_user():
    """
    Safely retrieve current user from session with fallback logic.
    Handles both session structures and prevents KeyError.
    
    Returns:
        dict: User data with email, first_name, last_name, role
        None: If no valid session exists
    """
    # Try nested structure first (from login route)
    user = session.get('user')
    if user and isinstance(user, dict):
        return user
    
    # Fallback: Try direct email key (legacy/old sessions)
    email = session.get('email')
    if email:
        return {'email': email}
    
    # No valid session
    return None

@app.route('/awaiting')
def awaiting():
    email = request.args.get('email', '')
    return render_template('awaiting.html', email=email)

@app.route('/api/auth/status')
def auth_status():
    email = request.args.get('email', '')
    if not email:
        return jsonify({"status": "error"}), 400
    user = users_collection.find_one({"email": email}, {"status": 1})
    if not user:
        return jsonify({"status": "error"}), 404
    return jsonify({"status": user.get("status", "pending")}), 200

@app.route('/forgot-password', methods=['GET'])
def forgot_password():
    return render_template('forgot_password.html')

@app.route('/api/auth/forgot-password', methods=['POST'])
def forgot_password_request():
    data  = request.get_json()
    email = data.get('email', '').strip()

    user = users_collection.find_one({"email": email})
    if not user:
        return jsonify({"status": "not_found"}), 200

    return jsonify({"status": "found"}), 200

@app.route('/api/auth/reset-password', methods=['POST'])
def self_reset_password():
    data             = request.get_json()
    email            = data.get('email', '').strip()
    new_password     = data.get('new_password', '')
    confirm_password = data.get('confirm_password', '')

    if new_password != confirm_password:
        return jsonify({"status": "error", "error": "Passwords do not match"}), 400

    if not re.match(r"(?=.*\d)(?=.*[a-z])(?=.*[A-Z]).{8,}", new_password):
        return jsonify({"status": "error", "error": "Password must be 8+ chars with uppercase, lowercase and number"}), 400

    user = users_collection.find_one({"email": email})
    if not user:
        return jsonify({"status": "error", "error": "Account not found"}), 404

    hashed = generate_password_hash(new_password, method='pbkdf2:sha256')
    users_collection.update_one(
        {"_id": user["_id"]},
        {"$set": {"password": hashed, "last_password_reset": datetime.now(NAIROBI_TZ)}}
    )

    db['session_logs'].insert_one({
        'email':      email,
        'action':     'password_reset_request',
        'reason':     'user_self_reset',
        'timestamp':  datetime.now(NAIROBI_TZ),
        'ip_address': request.remote_addr,
        'user_agent': request.headers.get('User-Agent')
    })

    return jsonify({"status": "success", "message": "Password updated. Please log in."}), 200

@app.route('/dashboard', methods=['GET'])
def dashboard():
   
    # 🆕 Use helper function for consistent session access
    current_user = get_current_user()
    
    if not current_user:
        flash('Please log in to continue.', 'warning')
        return redirect(url_for('auth'))
    
    user_email = current_user.get('email')
    
    # Fetch full user data from database
    user_data = db.users.find_one({"email": user_email})
    
    if not user_data:
        session.clear()
        flash('User account not found. Please log in again.', 'error')
        return redirect(url_for('auth'))
    
    # 🆕 Get last login (skip current session, get previous one)
    last_log = db.session_logs.find({
        "email": user_email, 
        "action": "login_success"
    }).sort("timestamp", -1).skip(1).limit(1)
    
    last_log_list = list(last_log)
    if last_log_list:
        last_login_timestamp = to_eat(last_log_list[0]['timestamp'])
    else:
        last_login_timestamp = None
    
    search_query = request.args.get('search', '').strip()
    
    # Set today's date range
    now = datetime.now(NAIROBI_TZ)
    today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    today_end = today_start + timedelta(days=1)
    
    # Fetch orders and retail collection
    all_orders = list(orders_collection.find().sort('date', -1))
    retail_collection = list(db["retail"].find({"date": now.strftime('%Y-%m-%d')}))
    
    # Calculate dashboard stats
    stats = calculate_dashboard_stats(all_orders, retail_collection, today_start, today_end)
    
    # Calculate additional metrics
    total_orders = len(all_orders)
    pending_count = sum(1 for order in all_orders if float(order.get('balance', 0)) > 0)
    completed_count = sum(1 for order in all_orders if float(order.get('balance', 0)) <= 0.001)
    
    # Today's debt (orders created today with balance)
    todays_debt = sum(
        float(order.get('balance', 0)) 
        for order in all_orders 
        if process_date(order.get('date')) >= today_start and 
           process_date(order.get('date')) < today_end and 
           float(order.get('balance', 0)) > 0
    )
    
    return render_template(
        'dashboard.html',
        user=user_data,
        now=now,
        last_login_timestamp=last_login_timestamp,
        total_sales_today=stats['total_sales_today'],
        retail_sales_today=stats['retail_sales_today'],
        wholesale_sales_today=stats['wholesale_sales_today'],
        total_debt=stats['total_debt'],  # 🆕 Changed from total_debts for consistency
        total_expenses=stats.get('total_expenses', 0),  # 🆕 Added
        open_orders_count=stats['open_orders_count'],
        closed_orders_count=stats['closed_orders_count'],
        retail_open_orders=stats['retail_open_orders'],
        retail_closed_orders=stats['retail_closed_orders'],
        wholesale_open_orders=stats['wholesale_open_orders'],
        wholesale_closed_orders=stats['wholesale_closed_orders'],
        todays_debt=todays_debt,
        search=search_query,
    )


@app.route('/api/dashboard/stats', methods=['GET'])
def get_dashboard_stats():
 
    try:
        current_user = get_current_user()
        
        if not current_user:
            return jsonify({"error": "Unauthorized"}), 401
        
        # Get today's date range
        now = datetime.now(NAIROBI_TZ)
        today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        today_end = today_start + timedelta(days=1)
        
        # Fetch orders and retail collection
        all_orders = list(orders_collection.find().sort('date', -1))
        retail_collection = list(db["retail"].find({"date": now.strftime('%Y-%m-%d')}))
        
        # Calculate stats
        stats = calculate_dashboard_stats(all_orders, retail_collection, today_start, today_end)
        
        # Calculate today's debt
        todays_debt = sum(
            float(order.get('balance', 0)) 
            for order in all_orders 
            if process_date(order.get('date')) >= today_start and 
               process_date(order.get('date')) < today_end and 
               float(order.get('balance', 0)) > 0
        )
        
        # Calculate time until midnight reset
        midnight = today_end
        time_until_reset = (midnight - now).total_seconds()
        hours_until_reset = int(time_until_reset // 3600)
        minutes_until_reset = int((time_until_reset % 3600) // 60)
        
        # Return comprehensive stats
        return jsonify({
            "status": "success",
            "timestamp": now.isoformat(),
            "session_info": {
                "session_started": session.get('login_time'),  # Add this
            },
            "stats": {
                "total_sales_today": stats['total_sales_today'],
                "retail_sales_today": stats['retail_sales_today'],
                "wholesale_sales_today": stats['wholesale_sales_today'],
                "total_debt": stats['total_debt'],
                "total_expenses": stats.get('total_expenses', 0),
                "open_orders_count": stats['open_orders_count'],
                "closed_orders_count": stats['closed_orders_count'],
                "retail_open_orders": stats['retail_open_orders'],
                "retail_closed_orders": stats['retail_closed_orders'],
                "wholesale_open_orders": stats['wholesale_open_orders'],
                "wholesale_closed_orders": stats['wholesale_closed_orders'],
                "todays_debt": todays_debt
            },
            "reset_info": {
                "next_reset": midnight.isoformat(),
                "hours_until_reset": hours_until_reset,
                "minutes_until_reset": minutes_until_reset
            }
        }), 200
        
    except Exception as e:
        logger.error(f"API stats error: {str(e)}")
        return jsonify({"error": "Failed to fetch stats"}), 500

@app.route('/api/activities', methods=['GET'])
@login_required
def get_activities():
    """
    Main API endpoint for activity handler
    Query params:
    - filter: all|pending|completed|expenses|gateway|modified|previous
    - date: YYYY-MM-DD (optional)
    - page: int (default 1)
    - per_page: int (default 50)
    """
    # Parse params
    filter_type = request.args.get('filter', 'all')
    selected_date = request.args.get('date', None)
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 50))
    
    # Get MongoDB collections
    db = get_db()
    orders_collection = db['orders']
    expenses_collection = db['expenses']
    
    # Determine mode
    mode = 'sorted' if selected_date else 'raw'
    
    # Parse date range if provided
    date_start = None
    date_end = None
    if selected_date:
        date_obj = datetime.strptime(selected_date, '%Y-%m-%d')
        date_start = NAIROBI_TZ.localize(date_obj.replace(hour=0, minute=0, second=0))
        date_end = NAIROBI_TZ.localize(date_obj.replace(hour=23, minute=59, second=59))
    
    # Fetch and process activities based on filter and mode
    if mode == 'raw':
        activities = fetch_raw_activities(
            orders_collection, 
            expenses_collection, 
            filter_type, 
            page, 
            per_page
        )
        counts = calculate_filter_counts(orders_collection, expenses_collection)
        
        return jsonify({
            'mode': 'raw',
            'activities': activities['data'],
            'pagination': activities['pagination'],
            'counts': counts
        })
    
    else:  # sorted mode
        activities = fetch_sorted_activities(
            orders_collection,
            expenses_collection,
            filter_type,
            date_start,
            date_end,
            page,
            per_page
        )
        stats = calculate_stats(activities['data'], filter_type)
        counts = calculate_filter_counts(orders_collection, expenses_collection)
        
        return jsonify({
            'mode': 'sorted',
            'activities': activities['data'],
            'stats': stats,
            'pagination': activities['pagination'],
            'counts': counts
        })

@app.route('/export/sales/<format>')
def export_sales_data(format):
    time_filter = request.args.get('time') or 'all'
    status_filter = request.args.get('status') or 'all'
    search = request.args.get('search') or ''
    
    query = {}
    expense_query = {} # Expenses usually don't have a 'status'

    # Time Filtering
    if time_filter != 'all':
        now = datetime.now(NAIROBI_TZ)
        if time_filter == 'day':
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_filter == 'week':
            start = now - timedelta(days=now.weekday())
            start = start.replace(hour=0, minute=0, second=0, microsecond=0)
        elif time_filter == 'month':
            start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        query['date'] = {'$gte': start}
        expense_query['date'] = {'$gte': start}

    # Order Status Filter
    if status_filter == 'pending':
        query['balance'] = {'$gt': 0}
    elif status_filter == 'completed':
        query['balance'] = {'$lte': 0}

    # Search logic
    if search:
        query['$or'] = [
            {'receipt_id': {'$regex': search, '$options': 'i'}},
            {'salesperson_name': {'$regex': search, '$options': 'i'}}
        ]
        expense_query['description'] = {'$regex': search, '$options': 'i'}

    orders = list(db["orders"].find(query).sort('date', -1))
    expenses = list(db["expenses"].find(expense_query).sort('date', -1))
    
    if format == 'pdf':
        return generate_pdf_export(orders, expenses, time_filter, status_filter)
    # Note: docx would need similar logic updates
    return jsonify({'error': 'PDF recommended for this layout'}), 400


# ORDERS MNGMNT

@app.route('/orders', methods=['GET', 'POST'])
def orders():
    logger.info("Processing /orders route")

    if request.method == 'POST':
        try:
            now = datetime.now(NAIROBI_TZ)
            shop_name = request.form.get('shop_name', 'Retail Direct')
            salesperson_name = f"{session['user']['first_name']} {session['user']['last_name']}"
            order_type = request.form.get('order_type', 'wholesale')
            change = float(request.form.get('change', '0') or 0)
            items_raw = request.form.getlist('items[]')

            # FIXED: CALCULATE total_amount FIRST
            items = []
            total_amount = 0

            for i in range(0, len(items_raw), 2):
                try:
                    product_data = items_raw[i].split('|')
                    if len(product_data) >= 6 and product_data[0] == 'product':
                        product_name = product_data[1]
                        qty_str = items_raw[i + 1] if i + 1 < len(items_raw) else '0'
                        quantity = float(qty_str) if qty_str.replace('.', '').replace('-', '').isdigit() else 0.0
                        price = float(product_data[5])
                        amount = quantity * price
                        if quantity > 0:
                            total_amount += amount
                            items.extend(['product', product_name, 'quantity', quantity, 'price', price])
                            stock = stock_collection.find_one({'stock_name': product_name})
                            if stock:
                                current_quantity = float(stock.get('stock_quantity', 0))
                                if current_quantity >= quantity:
                                    stock_collection.update_one(
                                        {'stock_name': product_name},
                                        {'$inc': {'stock_quantity': -quantity}}
                                    )
                                    log_stock_change(stock.get('category', 'Unknown'), product_name, 'order_reduction', -quantity, price)
                                else:
                                    return jsonify({'error': f"Insufficient stock for {product_name}"}), 400
                except (IndexError, ValueError) as e:
                    logger.error(f"Error processing item: {e}")
                    continue

            if not items:
                return jsonify({'error': 'No valid items in order'}), 400

            # NOW handle payments AFTER total_amount is known
            is_dual_payment = request.form.get('payment_type_dual') == 'true'
            if is_dual_payment:
                cash_amount = float(request.form.get('cash_amount', '0') or 0)
                mpesa_amount = float(request.form.get('mpesa_amount', '0') or 0)
                total_amount_paid = float(request.form.get('total_amount_paid', '0') or 0)
                
                if cash_amount + mpesa_amount != total_amount_paid:
                    logger.error(f"Dual payment amounts don't match: cash={cash_amount}, mpesa={mpesa_amount}, total={total_amount_paid}")
                    return jsonify({'error': 'Payment amounts do not match'}), 400
                
                payment_type = 'dual'
                amount_paid = total_amount_paid
                
                payment_history = []
                if cash_amount > 0:
                    payment_history.append({
                        'amount': cash_amount,
                        'date': datetime.now(NAIROBI_TZ),
                        'payment_type': 'cash'
                    })
                if mpesa_amount > 0:
                    payment_history.append({
                        'amount': mpesa_amount,
                        'date': datetime.now(NAIROBI_TZ),
                        'payment_type': 'mpesa'
                    })
                
                logger.info(f"Processing dual payment: cash={cash_amount}, mpesa={mpesa_amount}, total={total_amount_paid}")
            else:
                payment_type = request.form.get('payment_type', 'cash')
                amount_paid = float(request.form.get('amount_paid', '0') or 0)
                payment_history = [{
                    'amount': min(amount_paid, total_amount),  # FIXED: Now total_amount exists
                    'date': datetime.now(NAIROBI_TZ),
                    'payment_type': payment_type
                }] if amount_paid > 0 else []

            # Validate payment_type for restricted clientsValidate
            restricted_clients = ['client', 'clients', 'walk in', 'walkin']
            if payment_type == 'credit' and shop_name.lower() in restricted_clients:
                logger.error(f"Credit not allowed for client: {shop_name}")
                return jsonify({'error': 'Credit payment is not allowed for walk-in or unspecified clients'}), 400

            receipt_id = get_next_receipt_id()
            balance = max(total_amount - amount_paid, 0)

            # Build order data
            order_data = {
                'receipt_id': receipt_id,
                'salesperson_name': salesperson_name,
                'shop_name': shop_name,
                'salesperson_name_lower': salesperson_name.lower(),
                'shop_name_lower': shop_name.lower(),
                'items': items,
                'payment': min(amount_paid, total_amount),
                'balance': balance,
                'pending_payment': 0.0,
                'payment_history': payment_history,
                'date': datetime.now(NAIROBI_TZ),
                'order_type': order_type,
                'payment_type': payment_type,
                'change': change,
                'closed_date': datetime.now(NAIROBI_TZ) if balance == 0 else None,
                'tracking': {
                    'status': 'pending',
                    'last_updated': datetime.now(NAIROBI_TZ),
                    'notes': 'Order received, awaiting dispatch'
                }
            }

            if is_dual_payment:
                order_data['payment_breakdown'] = {
                    'cash': cash_amount,
                    'mpesa': mpesa_amount
                }

            orders_collection.insert_one(order_data)
            
            # Create notification for new order
            notification_id = f"ORD-{receipt_id}-{int(datetime.now(NAIROBI_TZ).timestamp())}"
            
            # Determine payment status display
            if balance == 0:
                status = 'paid'
                amount_display = total_amount
            else:
                status = 'credit'
                amount_display = f"{amount_paid}/{total_amount}"  # Shows 130/200 format
            
            notifications_collection.insert_one({
                'notification_id': notification_id,
                'receipt_id': receipt_id,
                'salesperson_name': f"{salesperson_name} created {order_type} order",  # Removed duplicate
                'shop_name': shop_name,
                'amount': amount_display,  # Now shows paid/total format
                'order_type': order_type,
                'type': 'order',
                'status': status,
                'created_at': datetime.now(NAIROBI_TZ),
                'read': False,
                'category': 'orders'
            })


            client = clients_collection.find_one({'shop_name': shop_name})
            if client:
                new_debt = client.get('debt', 0) + balance
                clients_collection.update_one(
                    {'shop_name': shop_name},
                    {'$set': {'debt': new_debt}}
                )
            else:
                clients_collection.insert_one({
                    'shop_name': shop_name,
                    'debt': balance,
                    'created_at': datetime.now(NAIROBI_TZ),
                    'location': None
                })

            log_user_action('Opened Order', f"Order #{receipt_id} - {order_type} for {shop_name}")
            logger.info(f"Order {receipt_id} created with {'dual payment: cash='+str(cash_amount)+', mpesa='+str(mpesa_amount) if is_dual_payment else f'single payment: {payment_type}={amount_paid}'}")
            return jsonify({'message': 'Order created successfully', 'receipt_id': receipt_id}), 200

        except Exception as e:
            logger.error(f"Error creating order: {e}")
            return jsonify({'error': str(e)}), 500

    # GET method
    try:
        orders = []
        for doc in orders_collection.find().sort('date', -1):
            try:  # ADD THIS
                order_dict = doc
                
            
                items_raw = order_dict.get('items', [])
                items_list = []
            
                if order_dict.get('order_type') == 'app' and isinstance(items_raw, list) and items_raw and isinstance(items_raw[0], dict):
                    for item in items_raw:
                        items_list.append({
                            'name': item.get('product', 'Unknown'),
                            'quantity': int(item.get('quantity', '0')),
                            'price': float(item.get('price', '0.0')),
                            'amount': int(item.get('quantity', '0')) * float(item.get('price', '0.0'))
                        })
                else:
                    i = 0
                    while i < len(items_raw):
                        if items_raw[i] == 'product':
                            product_name = items_raw[i + 1]
                            quantity_str = str(items_raw[i + 3]) if i + 2 < len(items_raw) and items_raw[i + 2] == 'quantity' else '0'
                            price_str = str(items_raw[i + 5]) if i + 4 < len(items_raw) and items_raw[i + 4] == 'price' else '0'
                            quantity = float(quantity_str) if quantity_str.replace('.', '').replace('-', '').isdigit() else 0.0
                            price = float(price_str) if price_str.replace('.', '').replace('-', '').isdigit() else 0.0
                            items_list.append({
                                'name': product_name,
                                'quantity': quantity,
                                'price': price,
                                'amount': quantity * price
                            })
                            i += 6
                        else:
                            i += 1
            
                order_response = {
                    'receipt_id': order_dict.get('receipt_id', str(doc['_id'])),
                    'salesperson_name': order_dict.get('salesperson_name', 'N/A'),
                    'salesperson_id': order_dict.get('salesperson_id', ''),
                    'shop_name': order_dict.get('shop_name', 'Unknown Shop'),
                    'total_items': process_items(order_dict.get('items')),
                    'items_list': items_list,
                    'payment': order_dict.get('payment', 0),
                    'balance': order_dict.get('balance', 0),
                    'date': process_date(order_dict.get('date')),
                    'closed_date': process_date(order_dict.get('closed_date', None)) if order_dict.get('closed_date') else None,
                    'order_type': order_dict.get('order_type', 'wholesale'),
                    'payment_type': order_dict.get('payment_type', 'cash'),
                    'change': order_dict.get('change', 0)
                }
            
                if 'payment_breakdown' in order_dict:
                    order_response['payment_breakdown'] = order_dict['payment_breakdown']
            
                orders.append(order_response)
        
            except Exception as e:  # MOVE THIS HERE
                print(f"Error processing order {doc.get('receipt_id')}: {e}")
                continue  # Skip this order, continue with next
    
        recent_activity = orders[:3]
        stock_items = list(stock_collection.find().sort('stock_name', 1))
        return render_template('orders.html', orders=orders, recent_activity=recent_activity, stock_items=stock_items, user=session['user'])

    except Exception as e:
        logger.error(f"Error fetching orders: {e}")
        return render_template('error.html', message=f"Failed to load orders: {str(e)}", user=session['user']), 500

@app.route('/orders_data', methods=['GET'])
def orders_data():
    """Return JSON data for orders, optionally filtered by shop_name."""
    shop_name = request.args.get('shop_name', '').strip()
    query = {'shop_name': shop_name} if shop_name else {}

    orders_list = []
    try:
        for doc in orders_collection.find(query).sort('date', -1):
            order_date = process_date(doc.get('date')) if doc.get('date') else None
            payment = float(doc.get('payment', 0))
            pending_payment = float(doc.get('pending_payment', 0))
            total_amount = 0

            try:
                items = doc.get('items', [])
                if isinstance(items, list) and items:
                    if doc.get('order_type') == 'app':
                        for item in items:
                            if isinstance(item, dict):
                                total_amount += float(item.get('price', 0)) * float(item.get('quantity', 0))
                    else:
                        for i in range(0, len(items), 6):
                            if i + 5 < len(items) and items[i] == 'product':
                                price = float(items[i + 5])
                                quantity = float(items[i + 3])
                                total_amount += price * quantity
                if doc.get('order_type') == 'app':
                    total_amount += float(doc.get('delivery_fee', 0))
            except (TypeError, ValueError, IndexError) as e:
                logging.error(f"Error calculating total_amount for order {doc.get('receipt_id')}: {e}")
                total_amount = payment + pending_payment

            balance = total_amount - payment

            orders_list.append({
                'receipt_id': doc.get('receipt_id', str(doc['_id'])),
                'shop_name': doc.get('shop_name'),
                'date': order_date.isoformat() if order_date else None,
                'payment': payment,
                'pending_payment': pending_payment,
                'balance': balance,
                'total_amount': total_amount,
                'order_type': doc.get('order_type', 'unknown'),
                'salesperson_name': doc.get('salesperson_name'),
                'notes': doc.get('notes'),
                'status': doc.get('status', 'unknown')
            })
    except Exception as e:
        logging.error(f"Error fetching orders data: {e}")
        return jsonify([]), 200

    return jsonify(orders_list), 200

## partial payments

@app.route('/mark_paid/<receipt_id>', methods=['POST'])
def mark_paid(receipt_id):
    try:
        # Find order by receipt_id
        order = orders_collection.find_one({'receipt_id': receipt_id})
        if not order:
            return jsonify({"error": f"Order with receipt_id {receipt_id} not found"}), 404
        
        # Extract order data
        current_payment = float(order.get('payment', 0))
        current_balance = float(order.get('balance', 0))
        amount_paid = float(request.form.get('amount_paid', 0))
        now = datetime.now(NAIROBI_TZ)
        
        # Validation
        if amount_paid <= 0:
            return jsonify({"error": "Payment amount must be greater than 0"}), 400
        if current_balance <= 0:
            return jsonify({"error": "Order is already fully paid"}), 400
        
        # Update payment and balance
        new_payment = current_payment + amount_paid
        new_balance = max(current_balance - amount_paid, 0)
        payment_history = order.get('payment_history', [])
        
        # Handle dual payment
        payment_type = request.form.get('payment_type', 'cash')
        is_dual = request.form.get('is_dual_payment', 'false') == 'true'
        
        # FIXED: Append two separate entries for dual payments
        if is_dual or payment_type == 'dual':
            cash_amount = float(request.form.get('cash_amount', 0))
            mpesa_amount = float(request.form.get('mpesa_amount', 0))
            
            # Append cash entry
            if cash_amount > 0:
                payment_history.append({
                    'amount': cash_amount,
                    'date': now,
                    'payment_type': 'cash'
                })
            
            # Append mpesa entry
            if mpesa_amount > 0:
                payment_history.append({
                    'amount': mpesa_amount,
                    'date': now,
                    'payment_type': 'mpesa'
                })
        else:
            # Single payment entry
            payment_history.append({
                'amount': amount_paid,
                'date': now,
                'payment_type': payment_type
            })
        
        # Prepare update data
        update_data = {
            'payment': new_payment,
            'balance': new_balance,
            'payment_history': payment_history,
            'last_activity_date': now
        }
        
        if new_balance == 0:
            update_data['closed_date'] = now
        
        # Update order in MongoDB
        orders_collection.update_one(
            {'_id': order['_id']},
            {'$set': update_data}
        )
        
        # Create notification for payment
        # Create notification for payment
        notification_id = f"PAY-{receipt_id}-{int(now.timestamp())}"
        if new_balance == 0:
            title = "Order Fully Paid"
            status_text = f"{order.get('salesperson_name', 'User')} marked order #{receipt_id} as fully paid"
        else:
            title = "Partial Payment"
            status_text = f"{order.get('salesperson_name', 'User')} paid KSh {amount_paid:.2f} on #{receipt_id}. Balance: KSh {new_balance:.2f}"

        notifications_collection.insert_one({
            'notification_id': notification_id,
            'receipt_id': receipt_id,
            'salesperson_name': status_text,  # Using this field for the message
            'shop_name': order.get('shop_name', 'Unknown'),
            'amount': amount_paid,
            'order_type': order.get('order_type', 'retail'),
            'type': 'payment',
            'status': 'paid' if new_balance == 0 else 'partial_payment',
            'created_at': now,
            'read': False,
            'category': 'payments'
        })
        # Update client debt
        shop_name = order.get('shop_name')
        if shop_name:
            client = db["clients"].find_one({'shop_name': shop_name})
            if client:
                new_client_debt = max(float(client.get('debt', 0)) - amount_paid, 0)
                db["clients"].update_one(
                    {'_id': client['_id']},
                    {'$set': {'debt': new_client_debt}}
                )
        
        # Create notification
        
        
        # Log the action
        try:
            log_user_action('Payment Processed', f"Order #{receipt_id}: KSh {amount_paid} paid, new balance KSh {new_balance}")
        except Exception as log_error:
            logger.error(f"Failed to log action: {log_error}")
        
        return jsonify({
            "success": True, 
            "message": "Payment processed successfully", 
            "new_balance": new_balance
        })
        
    except Exception as e:
        logger.error(f"Error in mark_paid: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/orders/<receipt_id>/can-edit', methods=['GET'])
def check_edit_eligibility(receipt_id):
    """
    Check if order can be edited based on:
    1. Balance > 0 (not fully paid)
    2. Order age < 72 hours
    3. Original items >= 5
    """
    order = orders_collection.find_one({"receipt_id": receipt_id})
    
    if not order:
        return jsonify({"can_edit": False, "reason": "Order not found"}), 404
    
    # CHECK 1: Balance must be > 0 (not fully paid)
    if order.get('balance', 0) <= 0:
        return jsonify({
            "can_edit": False, 
            "reason": "Cannot edit fully paid orders"
        }), 200
    
    # CHECK 2: Order age must be < 72 hours (3 days)
    order_date = order.get('date')
    if isinstance(order_date, str):
        order_date = datetime.fromisoformat(order_date.replace('Z', '+00:00'))
    
    now = datetime.now(NAIROBI_TZ)
    if order_date.tzinfo is None:
        order_date = NAIROBI_TZ.localize(order_date)
    
    age_hours = (now - order_date).total_seconds() / 3600
    
    if age_hours > 72:
        return jsonify({
            "can_edit": False, 
            "reason": f"Cannot edit orders older than 3 days (this order is {int(age_hours)} hours old)"
        }), 200
    
    # CHECK 3: Original items must be >= 5
    items = order.get('items', [])
    item_count = len([x for x in items if x == 'product'])
    
    if item_count < 5:
        return jsonify({
            "can_edit": False, 
            "reason": f"Cannot edit orders with less than 5 items (this order has {item_count} items)"
        }), 200
    
    # ALL CHECKS PASSED - Order can be edited
    # Parse items into structured format for frontend
    items_list = []
    i = 0
    while i < len(items):
        if items[i] == 'product':
            product_name = items[i + 1] if i + 1 < len(items) else 'Unknown'
            quantity = items[i + 3] if i + 3 < len(items) else 0
            price = items[i + 5] if i + 5 < len(items) else 0
            
            items_list.append({
                'product': product_name,
                'quantity': quantity,
                'price': price,
                'total': quantity * price
            })
            i += 6
        else:
            i += 1
    
    total = order.get('payment', 0) + order.get('balance', 0)
    
    return jsonify({
        "can_edit": True,
        "order": {
            "receipt_id": order['receipt_id'],
            "shop_name": order.get('shop_name', 'Unknown'),
            "salesperson_name": order.get('salesperson_name', 'N/A'),
            "order_type": order.get('order_type', 'wholesale'),
            "items": items_list,
            "item_count": item_count,
            "total": total,
            "payment": order.get('payment', 0),
            "balance": order.get('balance', 0),
            "order_date": order_date.isoformat(),
            "age_hours": int(age_hours)
        }
    }), 200


@app.route('/api/orders/<receipt_id>/edit', methods=['PUT'])
def edit_order(receipt_id):
    """
    Edit an order with constraints:
    - Can add unlimited items
    - Can remove max 2 items from original
    - Must keep at least 3 original items
    - Recalculate balance based on new total
    - Update client debt
    - Create notification
    """
    data = request.json
    
    # Get data from request
    items_to_keep = data.get('items_to_keep', [])  # Original items to keep
    items_removed = data.get('items_removed', [])  # Original items removed
    new_items = data.get('new_items', [])  # New items added
    edit_reason = data.get('reason', '').strip()
    edited_by = data.get('edited_by', 'Current User')
    
    # Validation
    if not edit_reason:
        return jsonify({"error": "Reason for edit is required"}), 400
    
    # Get current order
    order = orders_collection.find_one({"receipt_id": receipt_id})
    if not order:
        return jsonify({"error": "Order not found"}), 404
    
    # Re-check eligibility (in case state changed)
    if order.get('balance', 0) <= 0:
        return jsonify({"error": "Cannot edit fully paid orders"}), 403
    
    order_date = order.get('date')
    if isinstance(order_date, str):
        order_date = datetime.fromisoformat(order_date.replace('Z', '+00:00'))
    
    now = datetime.now(NAIROBI_TZ)
    if order_date.tzinfo is None:
        order_date = NAIROBI_TZ.localize(order_date)
    
    age_hours = (now - order_date).total_seconds() / 3600
    if age_hours > 72:
        return jsonify({"error": "Cannot edit orders older than 3 days"}), 403
    
    # Get original items
    original_items = order.get('items', [])
    original_item_count = len([x for x in original_items if x == 'product'])
    
    if original_item_count < 5:
        return jsonify({"error": "Cannot edit orders with less than 5 items"}), 403
    
    # CHECK REMOVAL LIMIT: Max 2 items can be removed
    if len(items_removed) > 2:
        return jsonify({
            "error": f"Cannot remove more than 2 items. You tried to remove {len(items_removed)} items."
        }), 400
    
    # CHECK MINIMUM ITEMS: Must keep at least 3 original items
    if len(items_to_keep) < 3:
        return jsonify({
            "error": f"Must keep at least 3 original items. You are keeping only {len(items_to_keep)} items."
        }), 400
    
    # PRESERVE ORIGINAL SNAPSHOT (if not already saved)
    if not order.get('original_snapshot'):
        original_total = order.get('payment', 0) + order.get('balance', 0)
        original_snapshot = {
            "items": original_items,
            "total": original_total,
            "payment": order.get('payment', 0),
            "balance": order.get('balance', 0),
            "date": order_date,
            "items_count": original_item_count
        }
    else:
        original_snapshot = order.get('original_snapshot')
    
    # CALCULATE NEW TOTALS
    
    # 1. Calculate kept items total
    kept_items_total = 0
    kept_items_flat = []
    for item in items_to_keep:
        product = item.get('product', '')
        quantity = float(item.get('quantity', 0))
        price = float(item.get('price', 0))
        total = quantity * price
        
        kept_items_flat.extend(['product', product, 'quantity', quantity, 'price', price])
        kept_items_total += total
    
    # 2. Calculate new items total
    new_items_total = 0
    new_items_flat = []
    for item in new_items:
        product = item.get('product', '')
        quantity = float(item.get('quantity', 0))
        price = float(item.get('price', 0))
        total = quantity * price
        
        new_items_flat.extend(['product', product, 'quantity', quantity, 'price', price])
        new_items_total += total
    
    # 3. Calculate removed items total
    removed_items_total = 0
    removed_items_names = []
    for item in items_removed:
        quantity = float(item.get('quantity', 0))
        price = float(item.get('price', 0))
        removed_items_total += (quantity * price)
        removed_items_names.append(item.get('product', 'Unknown'))
    
    # 4. New grand total
    new_grand_total = kept_items_total + new_items_total
    
    # 5. Calculate new balance
    amount_paid = order.get('payment', 0)
    new_balance = new_grand_total - amount_paid
    
    # 6. Get old values for comparison
    old_total = order.get('payment', 0) + order.get('balance', 0)
    old_balance = order.get('balance', 0)
    
    # COMBINE ALL ITEMS (kept + new)
    final_items_flat = kept_items_flat + new_items_flat
    final_item_count = len(items_to_keep) + len(new_items)
    
    # CREATE EDIT TAG
    added_items_names = [item.get('product', 'Unknown') for item in new_items]
    
    edit_tag = {
        "modification_date": now,
        "original_order_date": order_date.strftime('%d/%m/%Y %H:%M'),
        "old_items_count": original_item_count,
        "new_items_count": final_item_count,
        "items_kept": len(items_to_keep),
        "items_removed_count": len(items_removed),
        "items_added_count": len(new_items),
        "items_removed_list": ", ".join(removed_items_names) if removed_items_names else "None",
        "items_added_list": ", ".join(added_items_names) if added_items_names else "None",
        "old_total": old_total,
        "new_total": new_grand_total,
        "old_balance": old_balance,
        "new_balance": new_balance,
        "amount_paid": amount_paid,
        "price_change": new_grand_total - old_total,
        "balance_change": new_balance - old_balance,
        "edited_by": edited_by,
        "reason": edit_reason
    }
    
    # UPDATE ORDER IN DATABASE
    update_result = orders_collection.update_one(
        {"receipt_id": receipt_id, "balance": {"$gt": 0}},
        {
            "$set": {
                "items": final_items_flat,
                "balance": new_balance,
                "pending_payment": max(0, new_balance),
                "last_edited": now,
                "edit_tag": edit_tag,
                "original_snapshot": original_snapshot,
                "closed_date": now if new_balance <= 0 else None  # Close if overpaid
            },
            "$push": {
                "edit_history": {
                    "timestamp": now,
                    "edited_by": edited_by,
                    "reason": edit_reason,
                    "old_total": old_total,
                    "new_total": new_grand_total,
                    "old_balance": old_balance,
                    "new_balance": new_balance,
                    "items_removed": len(items_removed),
                    "items_added": len(new_items),
                    "changes": edit_tag
                }
            }
        }
    )
    
    if update_result.matched_count == 0:
        return jsonify({"error": "Order state changed, refresh and try again"}), 409
    
    # UPDATE CLIENT DEBT
    shop_name = order.get('shop_name')
    if shop_name:
        client = clients_collection.find_one({'shop_name': shop_name})
        if client:
            # Client's debt = Old debt - Old balance + New balance
            old_client_debt = float(client.get('debt', 0))
            new_client_debt = old_client_debt - old_balance + new_balance
            new_client_debt = max(0, new_client_debt)  # Can't be negative
            
            clients_collection.update_one(
                {'shop_name': shop_name},
                {'$set': {'debt': new_client_debt}}
            )
    
    # CREATE NOTIFICATION
    notification_id = f"EDIT-{receipt_id}-{int(now.timestamp())}"
    
    # Build change summary
    change_summary = []
    if len(new_items) > 0:
        change_summary.append(f"+{len(new_items)} item(s)")
    if len(items_removed) > 0:
        change_summary.append(f"-{len(items_removed)} item(s)")
    
    change_text = ", ".join(change_summary) if change_summary else "Modified"
    
    # Determine status
    if new_balance <= 0:
        status = 'paid'
        amount_display = new_grand_total
    else:
        status = 'credit'
        amount_display = f"{amount_paid}/{new_grand_total}"
    
    notifications_collection.insert_one({
        'notification_id': notification_id,
        'receipt_id': receipt_id,
        'salesperson_name': f"{edited_by} modified {order.get('order_type', 'retail')} order",
        'shop_name': shop_name,
        'amount': amount_display,
        'order_type': order.get('order_type', 'retail'),
        'type': 'order_edit',
        'status': status,
        'created_at': now,
        'read': False,
        'category': 'orders',
        'edit_details': {
            'changes': change_text,
            'old_total': old_total,
            'new_total': new_grand_total,
            'new_balance': new_balance
        }
    })
    
    # LOG ACTION
    try:
        log_user_action(
            'Order Modified', 
            f"Order #{receipt_id}: {change_text}, Old Total: KSh {old_total}, New Total: KSh {new_grand_total}, New Balance: KSh {new_balance}"
        )
    except:
        pass  # Don't fail if logging fails
    
    return jsonify({
        "success": True,
        "message": "Order updated successfully",
        "order": {
            "receipt_id": receipt_id,
            "new_total": new_grand_total,
            "new_balance": new_balance,
            "amount_paid": amount_paid,
            "items_count": final_item_count,
            "changes": {
                "added": len(new_items),
                "removed": len(items_removed),
                "kept": len(items_to_keep)
            }
        }
    }), 200

# STOCK MNGMNT
def create_stock_notification(action, stock_name, quantity_or_price, user):
    """Create notification for stock changes"""
    last_notif = notifications_collection.find_one(sort=[('notification_id', -1)])
    next_id = f"NOTIF{str(int(last_notif['notification_id'].replace('NOTIF', '')) + 1).zfill(6)}" if last_notif else 'NOTIF000001'
    
    messages = {
        'add_stock': f"New stock '{stock_name}' added ({quantity_or_price} units)",
        'restock': f"'{stock_name}' restocked (+{quantity_or_price} units)",
        'update_price': f"Price updated for '{stock_name}'",
        'update_price_and_category': f"Details updated for '{stock_name}'",
        'edit_stock_name': f"Stock renamed to '{stock_name}'"
    }
    
    notifications_collection.insert_one({
        'notification_id': next_id,
        'type': 'stock_change',
        'category': 'stock',
        'message': messages.get(action, f"Stock '{stock_name}' modified"),
        'stock_name': stock_name,
        'action': action,
        'user': user,
        'created_at': datetime.now(NAIROBI_TZ),
        'read': False
    })

@app.route('/stock', methods=['GET', 'POST'])
def stock():
    """Handle stock management with pagination."""
    cache_cleared = False
    client_logs = []

    if request.method == 'POST':
        if session['user']['role'] != 'manager':
            return jsonify({'status': 'error', 'error': 'Unauthorized: Only managers can modify stock'}), 403

        action = request.form.get('action')
        print(f"[STOCK_ROUTE] Processing action: {action}")
        client_logs.append(f"Processing action: {action}")

        if action == 'add_stock':
            print("[STOCK_ROUTE] Entering add_stock action")
            client_logs.append("Entering add_stock action")
            stock_name = request.form.get('stock_name')
            category = request.form.get('category')
            new_category = request.form.get('new_category')
            initial_quantity = request.form.get('initial_quantity')
            reorder_quantity = request.form.get('reorder_quantity')
            selling_price = request.form.get('selling_price')
            wholesale_price = request.form.get('wholesale_price')
            company_price = request.form.get('company_price')
            expire_date = request.form.get('expire_date')

            if not all([stock_name, category or new_category, initial_quantity, reorder_quantity, selling_price, wholesale_price, company_price, expire_date]):
                print("[STOCK_ROUTE] Error: Missing required fields in add_stock")
                client_logs.append("Error: Missing required fields in add_stock")
                create_stock_notification('add_stock', stock_name, initial_quantity, f"{session['user']['first_name']} {session['user']['last_name']}")
                return jsonify({'status': 'error', 'error': 'All fields are required'}), 400

            try:
                initial_quantity = int(initial_quantity)
                reorder_quantity = int(reorder_quantity)
                selling_price = float(selling_price)
                wholesale_price = float(wholesale_price)
                company_price = float(company_price)
                if any(x < 0 for x in [initial_quantity, reorder_quantity, selling_price, wholesale_price, company_price]):
                    print("[STOCK_ROUTE] Error: Negative values detected in add_stock")
                    client_logs.append("Error: Negative values detected in add_stock")
                    return jsonify({'status': 'error', 'error': 'Numeric fields cannot be negative'}), 400
                datetime.strptime(expire_date, '%Y-%m-%d')
            except ValueError:
                print("[STOCK_ROUTE] Error: Invalid numeric or date format in add_stock")
                client_logs.append("Error: Invalid numeric or date format in add_stock")
                return jsonify({'status': 'error', 'error': 'Invalid numeric or date format'}), 400

            final_category = new_category.strip() if new_category else category
            category_prefix = ''.join(c for c in final_category[:3] if c.isalnum()).upper()
            counter = metadata_collection.find_one_and_update(
                {'_id': 'stock_counter'},
                {'$inc': {'last_id': 1}},
                upsert=True,
                return_document=True
            )
            new_counter = counter.get('last_id', 1) if counter else 1
            stock_id = f"{category_prefix}{new_counter:03d}"

            if stock_collection.find_one({'stock_name': stock_name}):
                print(f"[STOCK_ROUTE] Error: Stock item '{stock_name}' already exists")
                client_logs.append(f"Error: Stock item '{stock_name}' already exists")
                return jsonify({'status': 'error', 'error': f"Stock item '{stock_name}' already exists"}), 400
            if stock_collection.find_one({'stock_id': stock_id}):
                print(f"[STOCK_ROUTE] Error: Stock ID '{stock_id}' already exists")
                client_logs.append(f"Error: Stock ID '{stock_id}' already exists")
                return jsonify({'status': 'error', 'error': f"Stock ID '{stock_id}' already exists"}), 400

            stock_data = {
                'id': new_counter,
                'stock_id': stock_id,
                'stock_name': stock_name,
                'stock_quantity': initial_quantity,
                'reorder_quantity': reorder_quantity,
                'supplier_id': None,
                'company_price': company_price,
                'selling_price': selling_price,
                'wholesale': wholesale_price,
                'barprice': 0.0,
                'category': final_category,
                'date': datetime.now(NAIROBI_TZ).strftime('%Y-%m-%d %H:%M:%S'),
                'expire_date': expire_date,
                'uom': None,
                'code': stock_id,
                'date2': None
            }

            doc_id = stock_id.replace('/', '-')
            stock_collection.insert_one(stock_data)
            log_stock_change(final_category, stock_name, 'add_stock', initial_quantity, selling_price)
            log_stock_change(final_category, stock_name, 'wholesale_price_set', 0, wholesale_price)
            cache_cleared = clear_stock_cache_logic()
            notification_id = f"STOCK-{stock_id}-{int(datetime.now(NAIROBI_TZ).timestamp())}"
            notifications_collection.insert_one({
                'notification_id': notification_id,
                'type': 'stock_change',
                'category': 'stock',
                'message': f"New stock '{stock_name}' added ({initial_quantity} units)",
                'stock_name': stock_name,
                'action': 'add_stock',
                'user': session['user']['email'],
                'created_at': datetime.now(NAIROBI_TZ),
                'read': False
            })
            print(f"[STOCK_ROUTE] add_stock completed, cache cleared: {cache_cleared}")
            return jsonify({'status': 'success', 'message': 'Stock added successfully'}), 200

        elif action == 'restock':
            print("[STOCK_ROUTE] Entering restock action")
            client_logs.append("Entering restock action")
            stock_id = request.form.get('stock_id')
            if stock_id:
                stock = stock_collection.find_one({'stock_id': stock_id})
                if stock:
                    try:
                        restock_qty = int(request.form.get('restock_quantity', 0))
                        if restock_qty <= 0:
                            print("[STOCK_ROUTE] Error: Restock quantity must be positive")
                            client_logs.append("Error: Restock quantity must be positive")
                            return jsonify({'status': 'error', 'error': 'Restock quantity must be positive'}), 400
                        current_qty = stock.get('stock_quantity', 0)
                        stock_collection.update_one(
                            {'stock_id': stock_id},
                            {'$set': {'stock_quantity': current_qty + restock_qty}}
                        )
                        log_stock_change(stock.get('category'), stock.get('stock_name'), 'restock', restock_qty, stock.get('selling_price'))
                        create_stock_notification('restock', stock.get('stock_name'), restock_qty, f"{session['user']['first_name']} {session['user']['last_name']}")
                        cache_cleared = clear_stock_cache_logic()
                        print(f"[STOCK_ROUTE] restock completed, cache cleared: {cache_cleared}")
                        client_logs.append(f"restock completed, cache cleared: {cache_cleared}")
                        return jsonify({'status': 'success', 'message': 'Stock restocked successfully'}), 200
                    except ValueError:
                        print("[STOCK_ROUTE] Error: Invalid restock quantity")
                        client_logs.append("Error: Invalid restock quantity")
                        return jsonify({'status': 'error', 'error': 'Invalid restock quantity'}), 400
                else:
                    print(f"[STOCK_ROUTE] Error: Stock ID '{stock_id}' not found")
                    client_logs.append(f"Error: Stock ID '{stock_id}' not found")
                    return jsonify({'status': 'error', 'error': f"Stock ID '{stock_id}' not found"}), 404

        elif action == 'update_price':
            print("[STOCK_ROUTE] Entering update_price action")
            client_logs.append("Entering update_price action")
            stock_id = request.form.get('stock_id')
            if stock_id:
                stock = stock_collection.find_one({'stock_id': stock_id})
                if stock:
                    try:
                        new_selling_price = float(request.form.get('new_selling_price', 0))
                        new_wholesale_price = float(request.form.get('new_wholesale_price', 0))
                        if new_selling_price < 0 or new_wholesale_price < 0:
                            print("[STOCK_ROUTE] Error: Negative prices detected in update_price")
                            client_logs.append("Error: Negative prices detected in update_price")
                            return jsonify({'status': 'error', 'error': 'Prices cannot be negative'}), 400
                        updates = {}
                        if new_selling_price > 0:
                            updates['selling_price'] = new_selling_price
                        if new_wholesale_price > 0:
                            updates['wholesale'] = new_wholesale_price
                        if updates:
                            stock_collection.update_one({'stock_id': stock_id}, {'$set': updates})
                            if new_selling_price > 0:
                                log_stock_change(stock.get('category'), stock.get('stock_name'), 'price_update', 0, new_selling_price)
                            if new_wholesale_price > 0:
                                log_stock_change(stock.get('category'), stock.get('stock_name'), 'wholesale_price_update', 0, new_wholesale_price)
                            cache_cleared = clear_stock_cache_logic()
                            create_stock_notification('update_price', stock.get('stock_name'), new_selling_price or new_wholesale_price, f"{session['user']['first_name']} {session['user']['last_name']}")
                            print(f"[STOCK_ROUTE] update_price completed, cache cleared: {cache_cleared}")
                            client_logs.append(f"update_price completed, cache cleared: {cache_cleared}")
                            return jsonify({'status': 'success', 'message': 'Prices updated successfully'}), 200
                        else:
                            print("[STOCK_ROUTE] Error: No valid prices provided for update_price")
                            client_logs.append("Error: No valid prices provided for update_price")
                            return jsonify({'status': 'error', 'error': 'No valid prices provided'}), 400
                    except ValueError:
                        print("[STOCK_ROUTE] Error: Invalid price format in update_price")
                        client_logs.append("Error: Invalid price format in update_price")
                        return jsonify({'status': 'error', 'error': 'Invalid price format'}), 400
                else:
                    print(f"[STOCK_ROUTE] Error: Stock ID '{stock_id}' not found")
                    client_logs.append(f"Error: Stock ID '{stock_id}' not found")
                    return jsonify({'status': 'error', 'error': f"Stock ID '{stock_id}' not found"}), 404

        elif action == 'edit_stock_name':
            print("[STOCK_ROUTE] Entering edit_stock_name action")
            client_logs.append("Entering edit_stock_name action")
            stock_id = request.form.get('stock_id')
            new_stock_name = request.form.get('new_stock_name')
            if not stock_id or not new_stock_name:
                print("[STOCK_ROUTE] Error: Missing stock_id or new_stock_name")
                client_logs.append("Error: Missing stock_id or new_stock_name")
                return jsonify({'status': 'error', 'error': 'Stock ID and new stock name are required'}), 400
            stock = stock_collection.find_one({'stock_id': stock_id})
            if stock:
                existing_stock = stock_collection.find_one({'stock_name': new_stock_name})
                if existing_stock and existing_stock['stock_id'] != stock_id:
                    print(f"[STOCK_ROUTE] Error: Stock name '{new_stock_name}' already exists")
                    client_logs.append(f"Error: Stock name '{new_stock_name}' already exists")
                    return jsonify({'status': 'error', 'error': f"Stock name '{new_stock_name}' already exists"}), 400
                stock_collection.update_one({'stock_id': stock_id}, {'$set': {'stock_name': new_stock_name}})
                log_stock_change(stock.get('category'), new_stock_name, 'name_update', 0, stock.get('selling_price'))
                create_stock_notification('edit_stock_name', new_stock_name, 0, f"{session['user']['first_name']} {session['user']['last_name']}")
                cache_cleared = clear_stock_cache_logic()
                print(f"[STOCK_ROUTE] edit_stock_name completed, cache cleared: {cache_cleared}")
                client_logs.append(f"edit_stock_name completed, cache cleared: {cache_cleared}")
                return jsonify({'status': 'success', 'message': 'Stock name updated successfully'}), 200
            else:
                print(f"[STOCK_ROUTE] Error: Stock ID '{stock_id}' not found")
                client_logs.append(f"Error: Stock ID '{stock_id}' not found")
                return jsonify({'status': 'error', 'error': f"Stock ID '{stock_id}' not found"}), 404

        elif action == 'update_price_and_category':
            print("[STOCK_ROUTE] Entering update_price_and_category action")
            client_logs.append("Entering update_price_and_category action")
            stock_id = request.form.get('stock_id')
            if stock_id:
                stock = stock_collection.find_one({'stock_id': stock_id})
                if stock:
                    try:
                        updates = {}
                        new_selling_price = request.form.get('new_selling_price')
                        new_wholesale_price = request.form.get('new_wholesale_price')
                        new_company_price = request.form.get('new_company_price')
                        new_category = request.form.get('new_category')
                        new_category_input = request.form.get('new_category_input')

                        if new_selling_price:
                            new_selling_price = float(new_selling_price)
                            if new_selling_price < 0:
                                print("[STOCK_ROUTE] Error: Negative selling price detected")
                                client_logs.append("Error: Negative selling price detected")
                                return jsonify({'status': 'error', 'error': 'Selling price cannot be negative'}), 400
                            updates['selling_price'] = new_selling_price
                        if new_wholesale_price:
                            new_wholesale_price = float(new_wholesale_price)
                            if new_wholesale_price < 0:
                                print("[STOCK_ROUTE] Error: Negative wholesale price detected")
                                client_logs.append("Error: Negative wholesale price detected")
                                return jsonify({'status': 'error', 'error': 'Wholesale price cannot be negative'}), 400
                            updates['wholesale'] = new_wholesale_price
                        if new_company_price:
                            new_company_price = float(new_company_price)
                            if new_company_price < 0:
                                print("[STOCK_ROUTE] Error: Negative company price detected")
                                client_logs.append("Error: Negative company price detected")
                                return jsonify({'status': 'error', 'error': 'Company price cannot be negative'}), 400
                            updates['company_price'] = new_company_price
                        if new_category == 'new' and new_category_input:
                            updates['category'] = new_category_input.strip()
                        elif new_category:
                            updates['category'] = new_category

                        if not updates:
                            print("[STOCK_ROUTE] Error: No valid fields provided for update_price_and_category")
                            client_logs.append("Error: No valid fields provided for update_price_and_category")
                            return jsonify({'status': 'error', 'error': 'At least one field must be updated'}), 400

                        stock_collection.update_one({'stock_id': stock_id}, {'$set': updates})
                        if 'selling_price' in updates:
                            log_stock_change(stock.get('category'), stock.get('stock_name'), 'price_update', 0, updates['selling_price'])
                        if 'wholesale' in updates:
                            log_stock_change(stock.get('category'), stock.get('stock_name'), 'wholesale_price_update', 0, updates['wholesale'])
                        if 'company_price' in updates:
                            log_stock_change(stock.get('category'), stock.get('stock_name'), 'company_price_update', 0, updates['company_price'])
                        if 'category' in updates:
                            log_stock_change(updates['category'], stock.get('stock_name'), 'category_update', 0, stock.get('selling_price'))
                        cache_cleared = clear_stock_cache_logic()
                        create_stock_notification('update_price_and_category', stock.get('stock_name'), 0, f"{session['user']['first_name']} {session['user']['last_name']}")
                        print(f"[STOCK_ROUTE] update_price_and_category completed, cache cleared: {cache_cleared}")
                        client_logs.append(f"update_price_and_category completed, cache cleared: {cache_cleared}")
                        return jsonify({'status': 'success', 'message': 'Price and category updated successfully'}), 200
                    except ValueError:
                        print("[STOCK_ROUTE] Error: Invalid numeric format in update_price_and_category")
                        client_logs.append("Error: Invalid numeric format in update_price_and_category")
                        return jsonify({'status': 'error', 'error': 'Invalid numeric format'}), 400
                else:
                    print(f"[STOCK_ROUTE] Error: Stock ID '{stock_id}' not found")
                    client_logs.append(f"Error: Stock ID '{stock_id}' not found")
                    return jsonify({'status': 'error', 'error': f"Stock ID '{stock_id}' not found"}), 404

        print(f"[STOCK_ROUTE] Invalid action: {action}")
        client_logs.append(f"Invalid action: {action}")
        return jsonify({'status': 'error', 'error': 'Invalid action'}), 400

    # === GET: Paginated stock ===
    PAGE_SIZE = 30
    page = max(1, int(request.args.get('page', 1)))
    search = request.args.get('search', '').strip()
    category = request.args.get('category', '').strip()

    # Build query
    query = {}
    if search:
        query['stock_name'] = {'$regex': search, '$options': 'i'}
    if category:
        query['category'] = category

    # Total count
    total = stock_collection.count_documents(query)
    total_pages = (total + PAGE_SIZE - 1) // PAGE_SIZE
    skip = (page - 1) * PAGE_SIZE

    # Fetch current page
    cursor = stock_collection.find(query).sort('stock_name', 1).skip(skip).limit(PAGE_SIZE)
    today = datetime.now().date()
    stock_items = []

    for doc in cursor:
        item = clean_doc(doc)
        item['id'] = str(doc['_id'])

        # === Calculate days_left ===
        expire_str = doc.get('expire_date')
        if expire_str and expire_str != "0000-00-00 00:00:00":
            try:
                expire_date = datetime.strptime(expire_str.split()[0], '%Y-%m-%d').date()
                item['days_left'] = (expire_date - today).days
            except:
                item['days_left'] = None
        else:
            item['days_left'] = None

        stock_items.append(item)

    # Remove duplicates by stock_name (within page)
    seen = set()
    unique_stock_items = []
    for item in stock_items:
        stock_name = item['stock_name']
        if stock_name not in seen:
            seen.add(stock_name)
            unique_stock_items.append(item)
    stock_items = unique_stock_items

    # Expiry notifications (only for current page)
    for item in stock_items:
        expire_date = item.get('expire_date')
        if expire_date and expire_date != "0000-00-00 00:00:00":
            try:
                days_left = item.get('days_left')
                if days_left is not None and days_left <= 30:
                    notification_message = f"Stock '{item['stock_name']}' is nearing expiry ({days_left} days left) on {expire_date}"
                    if not notifications_collection.find_one({'message': notification_message}):
                        notifications_collection.insert_one({
                            'recipient': session['user']['uid'],
                            'message': notification_message,
                            'timestamp': datetime.now(NAIROBI_TZ),
                            'order_id': None,
                            'read': False
                        })
            except ValueError:
                continue

    # Get all unique categories for the filter dropdown
    all_categories = stock_collection.distinct('category')

    recent_activity = [
        {
            'receipt_id': doc.get('receipt_id', str(doc['_id'])),
            'salesperson_name': doc.get('salesperson_name', 'N/A'),
            'shop_name': doc.get('shop_name', 'Unknown Shop'),
            'date': process_date(doc.get('date'))
        }
        for doc in orders_collection.find().sort('date', -1).limit(3)
    ]

    # Pagination flags
    has_prev = page > 1
    has_next = page < total_pages

    print(f"[STOCK_ROUTE] Rendering page {page}/{total_pages} with {len(stock_items)} items")
    client_logs.append(f"Rendering page {page}/{total_pages} with {len(stock_items)} items")

    return render_template(
        'stock.html',
        stock_items=stock_items,
        all_categories=all_categories,
        page=page,
        total_pages=total_pages,
        has_prev=has_prev,
        has_next=has_next,
        search=search,
        category=category,
        recent_activity=recent_activity,
        client_logs=client_logs
    )

def receipts():
    try:
        page = request.args.get('page', 1, type=int)
        per_page = 30
        search = request.args.get('search', '').lower()
        period = request.args.get('period', 'all')
        
        # Build query
        query = {}
        if search:
            query["$or"] = [
                {"receipt_id": {"$regex": search, "$options": "i"}},
                {"shop_name": {"$regex": search, "$options": "i"}},
                {"salesperson_name": {"$regex": search, "$options": "i"}}
            ]
        
        if period != 'all':
            now = datetime.now(NAIROBI_TZ)
            if period == 'day':
                start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == 'week':
                start = now - timedelta(days=now.weekday())
                start = start.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == 'month':
                start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            query["date"] = {"$gte": start}
        
        # Get ALL matching orders first (for total count)
        all_orders = list(orders_collection.find(query).sort('date', -1))
        total_orders = len(all_orders)
        
        # Serialize ALL
        all_orders = [serialize_mongo_doc(order) for order in all_orders]
        
        # Pagination
        start_idx = (page - 1) * per_page
        paginated_orders = all_orders[start_idx:start_idx + per_page]
        
        return render_template('receipts.html', 
                             orders=paginated_orders, 
                             total_orders=total_orders,
                             user=session['user'])
    except Exception as e:
        logger.error(f"Error loading receipts: {str(e)}")
        return render_template('error.html', message=f"Error loading receipts: {str(e)}"), 500



@app.route('/clear_stock_cache', methods=['POST'])
def clear_stock_cache():
    """Clear stock cache with detailed logging."""
    print(f"[CLEAR_STOCK_CACHE] Request received: user={session['user']['email']}")
    try:
        print("[CLEAR_STOCK_CACHE] Calling update_stock_version")
        new_version, client_log_messages = update_stock_version()
        print("[CLEAR_STOCK_CACHE] Clearing stock_cache")
        stock_cache['data'] = None
        stock_cache['version'] = None
        stock_cache['timestamp'] = None
        print("[CLEAR_STOCK_CACHE] Stock cache cleared successfully")
        return jsonify({'status': 'success', 'message': 'Cache cleared', 'logs': client_log_messages}), 200
    except Exception as e:
        print(f"[CLEAR_STOCK_CACHE] Error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({'error': f'Failed to clear cache: {str(e)}', 'logs': client_log_messages}, session={'user': mock_user}), 500

@app.route('/stock_data', methods=['GET'])
def stock_data():
    """Fetch stock data or version from MongoDB for retail/wholesale modals."""
    try:
        version_doc = metadata_collection.find_one({'_id': 'stock_version'})
        current_version = str(version_doc.get('version', 0)) if version_doc else '0'

        if request.args.get('version_only') == 'true':
            return jsonify({'version': current_version}), 200

        if (stock_cache['data'] is not None and
                stock_cache['version'] == current_version and
                stock_cache['timestamp'] is not None and
                datetime.now() < stock_cache['timestamp'] + stock_cache['timeout']):
            print(f"Serving {len(stock_cache['data'])} stock items from cache (version: {current_version})")
            return jsonify({'version': current_version, 'data': stock_cache['data']}), 200

        stock_items = [
            {
                'stock_name': doc['stock_name'],
                'selling_price': float(doc.get('selling_price', 0)),
                'wholesale': float(doc.get('wholesale', 0)),
                'stock_quantity': float(doc.get('stock_quantity', 0)),
                'uom': doc.get('uom', 'Unit'),
                'category': doc.get('category', ''),
                'id': str(doc['_id']),
                'company_price': float(doc.get('company_price', 0)),
                'expire_date': doc.get('expire_date', None),
                'reorder_quantity': int(doc.get('reorder_quantity', 0))
            }
            for doc in stock_collection.find().sort('stock_name', 1)
        ]

        seen = set()
        unique_stock_items = []
        for item in stock_items:
            stock_name = item['stock_name']
            if stock_name not in seen and all(
                item[key] is not None for key in ['selling_price', 'wholesale', 'stock_quantity']
            ):
                seen.add(stock_name)
                unique_stock_items.append(item)

        if not unique_stock_items:
            print("No stock items found in MongoDB")
            return jsonify({'version': current_version, 'data': []}), 200

        stock_cache['data'] = unique_stock_items
        stock_cache['version'] = current_version
        stock_cache['timestamp'] = datetime.now()

        print(f"Returning {len(unique_stock_items)} stock items from MongoDB (version: {current_version})")
        return jsonify({'version': current_version, 'data': unique_stock_items}), 200
    except Exception as e:
        print(f"Error fetching stock data: {str(e)}")
        return jsonify({'error': 'Internal server error'}), 500
    

#        RECEIPTS

@app.route('/receipts')
def receipts():
    try:
        page = request.args.get('page', 1, type=int)
        per_page = 30
        search = request.args.get('search', '').lower()
        period = request.args.get('period', 'all')
        
        # Build query
        query = {}
        if search:
            query["$or"] = [
                {"receipt_id": {"$regex": search, "$options": "i"}},
                {"shop_name": {"$regex": search, "$options": "i"}},
                {"salesperson_name": {"$regex": search, "$options": "i"}}
            ]
        
        if period != 'all':
            now = datetime.now(NAIROBI_TZ)
            if period == 'day':
                start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == 'week':
                start = now - timedelta(days=now.weekday())
                start = start.replace(hour=0, minute=0, second=0, microsecond=0)
            elif period == 'month':
                start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            query["date"] = {"$gte": start}
        
        # Get ALL matching orders first (for total count)
        all_orders = list(orders_collection.find(query).sort('date', -1))
        total_orders = len(all_orders)
        
        # Serialize ALL
        all_orders = [serialize_mongo_doc(order) for order in all_orders]
        
        # Pagination
        start_idx = (page - 1) * per_page
        paginated_orders = all_orders[start_idx:start_idx + per_page]
        
        return render_template('receipts.html', 
                             orders=paginated_orders, 
                             total_orders=total_orders,
                             user=session['user'])
    except Exception as e:
        logger.error(f"Error loading receipts: {str(e)}")
        return render_template('error.html', message=f"Error loading receipts: {str(e)}"), 500

@app.route('/receipt/<order_id>')
def receipt(order_id):
    try:
        # ✅ FIX: Convert ObjectId if needed
        if ObjectId.is_valid(order_id):
            order_id_obj = ObjectId(order_id)
        else:
            order_id_obj = order_id
        order = orders_collection.find_one({'receipt_id': order_id})
        if not order:
            return render_template('error.html', message="Order not found"), 404
            
        # ✅ FIX: Serialize main order
        order = serialize_mongo_doc(order)
        
        items_raw = order.get('items', [])
        items_list = []
        subtotal_amount = 0

        if order.get('order_type') == 'app' and isinstance(items_raw, list) and items_raw and isinstance(items_raw[0], dict):
            for item in items_raw:
                quantity = float(item.get('quantity', '0'))
                price = float(item.get('price', '0.0'))
                amount = quantity * price
                subtotal_amount += amount
                items_list.append({
                    'name': item.get('product', 'Unknown'),
                    'quantity': quantity,
                    'price': price,
                    'amount': amount
                })
        else:
            i = 0
            while i < len(items_raw):
                if items_raw[i] == 'product':
                    product_name = items_raw[i + 1]
                    quantity_str = str(items_raw[i + 3]) if i + 2 < len(items_raw) and items_raw[i + 2] == 'quantity' else '0'
                    price_str = str(items_raw[i + 5]) if i + 4 < len(items_raw) and items_raw[i + 4] == 'price' else '0'
                    try:
                        quantity = float(quantity_str) if quantity_str.replace('.', '').replace('-', '').isdigit() else 0.0
                    except ValueError:
                        logger.error(f"Invalid quantity format for {product_name}: {quantity_str}")
                        quantity = 0.0
                    price = float(price_str) if price_str.replace('.', '').replace('-', '').isdigit() else 0.0
                    amount = quantity * price
                    subtotal_amount += amount
                    items_list.append({
                        'name': product_name,
                        'quantity': quantity,
                        'price': price,
                        'amount': amount
                    })
                    i += 6
                else:
                    i += 1

        shop_name = order.get('shop_name', 'Unknown Shop')
        try:
            shop = clients_collection.find_one({'name': shop_name})
            shop_address = shop.get('address', 'No address') if shop else 'No address'
        except Exception as e:
            logger.error(f"Error fetching shop address: {str(e)}")
            shop_address = 'No address available'

        order_data = {
            'receipt_id': order.get('receipt_id', order_id),
            'salesperson_name': order.get('salesperson_name', 'N/A'),
            'shop_name': shop_name,
            'shop_address': shop_address,
            'items_list': items_list,
            'total_items': process_items(order.get('items')),
            'subtotal': subtotal_amount,
            'total_amount': subtotal_amount,
            'payment': order.get('payment', 0),
            'balance': order.get('balance', 0),
            'date': order.get('date'),  # Already processed
            'order_type': order.get('order_type', 'wholesale'),
            'notes': order.get('notes', ''),
            'payment_history': order.get('payment_history', []),
            'status': order.get('status', 'pending')
        }
        
        # ✅ FIX: Serialize recent_activity
        recent_docs = list(orders_collection.find().sort('date', -1).limit(3))
        recent_activity = [
            {
                'receipt_id': doc.get('receipt_id', str(doc.get('_id', 'N/A'))),
                'salesperson_name': doc.get('salesperson_name', 'N/A'),
                'shop_name': doc.get('shop_name', 'Unknown Shop'),
                'date': process_date(doc.get('date'))
            }
            for doc in recent_docs
        ]
        
        logger.info(f"Order data for {order_id}: {order_data}")
        return render_template('receipt.html', order=order_data, recent_activity=recent_activity, user=session['user'])
    except Exception as e:
        logger.error(f"Error in receipt route for {order_id}: {str(e)}")
        return render_template('error.html', message=f"Internal Server Error: {str(e)}"), 500

@app.route('/api/orders/<receipt_id>', methods=['GET'])
def get_order(receipt_id):
    order = orders_collection.find_one({"receipt_id": receipt_id})
    
    if not order:
        return jsonify({"error": "Order not found"}), 404
    
    # Convert ObjectId to string
    order['_id'] = str(order['_id'])
    
    return jsonify(order), 200
        
#         CLIENTS

@app.route('/clients', methods=['GET'])
def clients():
    """Render the clients page with paginated data."""
    search_query = request.args.get('search', '').strip()
    page = int(request.args.get('page', 1))
    per_page = 12
    offset = (page - 1) * per_page

    # Initialize MongoDB query
    query = {}
    if search_query:
        query['shop_name_lower'] = {
            '$gte': search_query.lower(),
            '$lte': search_query.lower() + '\uf8ff'
        }

    # Fetch total count
    counter_doc = metadata_collection.find_one({'_id': 'clients_counter'})
    if counter_doc and not search_query:
        total_clients = counter_doc.get('count', 0)
    else:
        total_clients = clients_collection.count_documents(query)

    # Paginate query
    clients_cursor = clients_collection.find(query).sort('created_at', -1).skip(offset).limit(per_page)
    clients_list = []

    for doc in clients_cursor:
        shop_name = doc.get('shop_name', 'Unknown Shop')
        # Fetch latest order
        latest_order = orders_collection.find_one(
            {'shop_name': shop_name},
            sort=[('date', -1)]
        )
        last_order_date = None
        recent_order_amount = None
        if latest_order:
            try:
                last_order_date = process_date(latest_order.get('date'))
                items = latest_order.get('items', [])
                recent_order_amount = sum(
                    float(item[5]) * float(item[3])
                    for item in items
                    if isinstance(item, (list, tuple)) and len(item) > 5 and item[0] == 'product'
                )
            except (TypeError, IndexError, ValueError) as e:
                logging.error(f"Error calculating recent_order_amount for shop {shop_name}: {e}")
                recent_order_amount = 0.0

        clients_list.append({
            'shop_name': shop_name,
            'debt': float(doc.get('debt', 0)),
            'last_order_date': last_order_date,
            'recent_order_amount': recent_order_amount,
            'phone': doc.get('phone'),
            'location': doc.get('location'),
            'created_at': process_date(doc.get('created_at')),
            'order_types': doc.get('order_types', [])
        })

    total_pages = max(1, (total_clients + per_page - 1) // per_page)
    clients_with_debt = clients_collection.count_documents({'debt': {'$gt': 0}})

    return render_template(
        'clients.html',
        clients=clients_list,
        search=search_query,
        total_clients=total_clients,
        clients_with_debt=clients_with_debt,
        pagination={'page': page, 'per_page': per_page, 'total_pages': total_pages},
        user=session['user']
    )


@app.route('/clients_data', methods=['GET'])
def clients_data():
    """Return JSON data for clients with search filtering."""
    search_query = request.args.get('search', '').lower()
    query = {}
    if search_query:
        query['shop_name_lower'] = {
            '$gte': search_query,
            '$lte': search_query + '\uf8ff'
        }

    clients_list = []
    try:
        for doc in clients_collection.find(query).sort('created_at', -1):
            shop_name = doc.get('shop_name', 'Unknown Shop')
            try:
                # Ensure shop_name_lower exists
                if 'shop_name_lower' not in doc:
                    clients_collection.update_one(
                        {'_id': doc['_id']},
                        {'$set': {'shop_name_lower': shop_name.lower()}}
                    )

                created_at = process_date(doc.get('created_at')) if doc.get('created_at') else None
                last_order_date = process_date(doc.get('last_order_date')) if doc.get('last_order_date') else None

                clients_list.append({
                    'shop_name': shop_name,
                    'debt': float(doc.get('debt', 0)),
                    'last_order_date': last_order_date.isoformat() if last_order_date else None,
                    'recent_order_amount': float(doc.get('recent_order_amount', 0)),
                    'phone': doc.get('phone'),
                    'location': doc.get('location'),
                    'created_at': created_at.isoformat() if created_at else None,
                    'order_types': doc.get('order_types', [])
                })
            except Exception as e:
                logging.error(f"Error processing client {shop_name}: {e}")
                continue
    except Exception as e:
        logging.error(f"Error fetching clients data: {e}")
        return jsonify([]), 200

    return jsonify(clients_list), 200

@app.route('/edit_client/<shop_name>', methods=['POST'])
def edit_client(shop_name):
    """Edit an existing client’s details."""
    if session['user']['role'] != 'manager':
        return jsonify({'error': 'Unauthorized: Only managers can edit clients'}), 403

    original_shop_name = request.form.get('original_shop_name')
    new_shop_name = request.form.get('shop_name')
    phone = request.form.get('phone', None)
    location = request.form.get('location', None)

    if not new_shop_name:
        return jsonify({'error': 'Client name is required'}), 400

    try:
        client = clients_collection.find_one({'shop_name': original_shop_name})
        if not client:
            return jsonify({'error': 'Client not found'}), 404

        update_data = {}
        if new_shop_name != original_shop_name:
            update_data['shop_name'] = new_shop_name
            update_data['shop_name_lower'] = new_shop_name.lower()
        if phone is not None:
            update_data['phone'] = phone if phone else None
        if location is not None:
            update_data['location'] = location if location else None

        if update_data:
            clients_collection.update_one(
                {'shop_name': original_shop_name},
                {'$set': update_data}
            )
            if new_shop_name != original_shop_name:
                orders_collection.update_many(
                    {'shop_name': original_shop_name},
                    {'$set': {'shop_name': new_shop_name}}
                )
            log_user_action(
                'Edited Client',
                f"Updated client {original_shop_name} to {new_shop_name} - Phone: {phone}, Location: {location}"
            )

        return jsonify({'status': 'success', 'message': 'Client updated successfully'}), 200
    except Exception as e:
        logging.error(f"Error updating client {shop_name}: {e}")
        return jsonify({'error': 'Internal server error'}), 500


# LOADING SHEETS

@app.route('/load_to_loading_sheet/<receipt_id>/<action>')
def load_to_loading_sheet(receipt_id, action):
    # Find order in MongoDB
    order = db["orders"].find_one({'receipt_id': receipt_id})
    if not order:
        return "Order not found", 404
    
    if order.get('order_type', 'wholesale') == 'retail':
        return "Retail orders cannot be loaded to a loading sheet", 400
    
    # Parse items from MongoDB format
    items_raw = order.get('items', [])
    items_list = []
    i = 0
    while i < len(items_raw):
        if items_raw[i] == 'product':
            product_name = items_raw[i + 1]
            quantity = items_raw[i + 3] if i + 2 < len(items_raw) and items_raw[i + 2] == 'quantity' else 0
            items_list.append({'name': product_name, 'quantity': quantity})
            i += 6
        else:
            i += 1

    # If action is 'new', save the current sheet to MongoDB before creating a new one
    if action == 'new' and 'current_loading_sheet' in session:
        current_sheet = session['current_loading_sheet']
        items = current_sheet.get('items', [])
        total_items = current_sheet.get('total_items', 0)
        
        # Parse created_at from session
        created_at_str = current_sheet.get('created_at')
        if isinstance(created_at_str, str):
            try:
                created_at = datetime.fromisoformat(created_at_str)
            except ValueError:
                created_at = datetime.now(NAIROBI_TZ)
        else:
            created_at = datetime.now(NAIROBI_TZ)

        # Generate a unique loading sheet ID
        loading_sheet_id = f"LOAD_{datetime.now(NAIROBI_TZ).strftime('%Y%m%d_%H%M%S')}"
        
        # Save the current sheet to MongoDB
        db["loading_sheets"].insert_one({
            'loading_sheet_id': loading_sheet_id,
            'items': items,
            'total_items': total_items,
            'created_at': created_at,
            'user_id': session['user']['_id']
        })
        
        # Log the action
        log_user_action('Saved Loading Sheet', f"Saved loading sheet {loading_sheet_id} with {total_items} items")
        
        # Clear the current sheet from session
        session.pop('current_loading_sheet')

    # Now handle the new items
    if action == 'current' and 'current_loading_sheet' in session:
        current_items = session.get('current_loading_sheet', {}).get('items', [])
        for item in items_list:
            found = False
            for existing_item in current_items:
                if existing_item['name'] == item['name']:
                    existing_item['quantity'] += item['quantity']
                    found = True
                    break
            if not found:
                current_items.append(item)
        session['current_loading_sheet'] = {
            'items': current_items,
            'total_items': sum(item['quantity'] for item in current_items),
            'created_at': session.get('current_loading_sheet', {}).get('created_at', datetime.now(NAIROBI_TZ).isoformat())
        }
    else:
        # Create a new loading sheet in session
        session['current_loading_sheet'] = {
            'items': items_list,
            'total_items': sum(item['quantity'] for item in items_list),
            'created_at': datetime.now(NAIROBI_TZ).isoformat()
        }

    return redirect(url_for('loading_sheets'))

@app.route('/loading-sheets')
def loading_sheets():
    """Display the loading sheets page."""
    current_loading_sheet = session.get('current_loading_sheet', None)
    if current_loading_sheet:
        aggregated_items = current_loading_sheet.get('items', [])
        total_items = current_loading_sheet.get('total_items', 0)
        # Fix the datetime serialization issue
        created_at_str = current_loading_sheet.get('created_at')
        if isinstance(created_at_str, str):
            try:
                created_at = datetime.fromisoformat(created_at_str)
            except ValueError:
                created_at = datetime.now(NAIROBI_TZ)
        else:
            created_at = current_loading_sheet.get('created_at', datetime.now(NAIROBI_TZ))
    else:
        aggregated_items = []
        total_items = 0
        created_at = None

    # Get recent sheets from MongoDB
    try:
        recent_sheets = []
        for doc in db["loading_sheets"].find().sort('created_at', -1).limit(5):
            sheet_data = doc
            # Convert datetime if needed (MongoDB stores as datetime)
            created_at_field = sheet_data.get('created_at')
            if isinstance(created_at_field, str):
                try:
                    sheet_data['created_at'] = datetime.fromisoformat(created_at_field)
                except ValueError:
                    sheet_data['created_at'] = datetime.now(NAIROBI_TZ)
            # Already datetime, no change needed
            recent_sheets.append(sheet_data)
    except Exception as e:
        print(f"Error fetching recent sheets: {e}")
        recent_sheets = []

    now = datetime.now(NAIROBI_TZ)
    
    return render_template('maintainance.html', 
                          aggregated_items=aggregated_items, 
                          current_date=now, 
                          total_items=total_items, 
                          created_at=created_at, 
                          recent_sheets=recent_sheets)

@app.route('/view-loading-sheet')
def view_loading_sheet():
    """View a specific loading sheet."""
    sheet_id = request.args.get('sheet_id')
    print_mode = request.args.get('print') == 'true'
    
    if not sheet_id:
        flash('Sheet ID is required', 'error')
        return redirect(url_for('loading_sheets'))

    # Fetch loading sheet from MongoDB
    try:
        sheet = db["loading_sheets"].find_one({'loading_sheet_id': sheet_id})
        if not sheet:
            flash('Loading sheet not found', 'error')
            return redirect(url_for('loading_sheets'))
        
        # Handle date conversion
        created_at_field = sheet.get('created_at')
        if isinstance(created_at_field, str):
            try:
                created_at = datetime.fromisoformat(created_at_field)
            except ValueError:
                created_at = datetime.now(NAIROBI_TZ)
        else:
            created_at = created_at_field or datetime.now(NAIROBI_TZ)
        
        aggregated_items = sheet.get('items', [])
        total_items = sheet.get('total_items', 0)
        
        return render_template('view_loading_sheet.html',
                              aggregated_items=aggregated_items,
                              total_items=total_items,
                              created_at=created_at,
                              current_date=datetime.now(NAIROBI_TZ),
                              sheet_id=sheet_id,
                              print_mode=print_mode)
    except Exception as e:
        print(f"Error in view-loading-sheet: {str(e)}")
        flash(f'Error loading sheet: {str(e)}', 'error')
        return redirect(url_for('loading_sheets'))

@app.route('/download-loading-sheet')
def download_loading_sheet():
    sheet_id = request.args.get('sheet_id')
    
    # Handle specific sheet download if ID is provided
    if sheet_id:
        try:
            sheet = db["loading_sheets"].find_one({'loading_sheet_id': sheet_id})
            if not sheet:
                return "Loading sheet not found", 404
                
            aggregated_items = sheet.get('items', [])
            total_items = sheet.get('total_items', 0)
            
            # Handle created_at datetime conversion
            created_at_field = sheet.get('created_at')
            if isinstance(created_at_field, str):
                try:
                    created_at = datetime.fromisoformat(created_at_field)
                except ValueError:
                    created_at = datetime.now(NAIROBI_TZ)
            else:
                created_at = created_at_field or datetime.now(NAIROBI_TZ)
        except Exception as e:
            print(f"Error fetching loading sheet: {str(e)}")
            return f"Error fetching loading sheet: {str(e)}", 500
    # Handle current sheet in session
    else:
        current_loading_sheet = session.get('current_loading_sheet', None)
        if not current_loading_sheet or not current_loading_sheet.get('items'):
            return "No loading sheet available to download", 400

        aggregated_items = current_loading_sheet.get('items', [])
        total_items = current_loading_sheet.get('total_items', 0)
        
        # Fix created_at datetime handling
        created_at_str = current_loading_sheet.get('created_at')
        if isinstance(created_at_str, str):
            try:
                created_at = datetime.fromisoformat(created_at_str)
            except ValueError:
                created_at = datetime.now(NAIROBI_TZ)
        else:
            created_at = current_loading_sheet.get('created_at', datetime.now(NAIROBI_TZ))

    try:
        # Generate PDF
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4

        # Header
        p.setFont("Helvetica-Bold", 14)
        p.drawCentredString(width / 2, height - 50, "Dreamland Distributors")
        p.setFont("Helvetica", 10)
        p.drawCentredString(width / 2, height - 70, "P.O Box 123-00200 Nairobi | Phone: 0725 530632")
        p.line(50, height - 80, width - 50, height - 80)
        p.setFont("Helvetica-Bold", 12)
        p.drawCentredString(width / 2, height - 100, "Loading Sheet")
        p.setFont("Helvetica", 10)
        formatted_date = created_at.strftime('%d/%m/%Y %H:%M')
        p.drawString(50, height - 120, f"Date: {formatted_date}")
        
        # Sheet ID
        sheet_name = sheet_id if sheet_id else "Current"
        p.drawString(450, height - 120, f"Sheet: {sheet_name}")

        # Table header
        y = height - 160
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, y, "Item")
        p.drawString(300, y, "Details")
        p.line(50, y - 5, width - 50, y - 5)
        y -= 20

        # Items
        p.setFont("Helvetica", 9)
        for item in aggregated_items:
            if y < 100:  # Start new page if not enough space
                p.showPage()
                p.setFont("Helvetica", 9)
                y = height - 50
            
            # Truncate long item names
            item_name = item['name'][:35] + "..." if len(item['name']) > 35 else item['name']
            
            p.drawString(50, y, item_name)
            
            if "sugar" in item['name'].lower() and "2k" in item['name'].lower():
                notes = f"2 pieces x {item['quantity']}"
            elif item['quantity'] > 1:
                notes = f"{item['quantity']} pieces"
            else:
                notes = "Single unit"
                
            p.drawString(300, y, notes)
            y -= 18

        # Footer
        y -= 20
        p.line(50, y, width - 50, y)
        y -= 20
        
        p.setFont("Helvetica-Bold", 12)
        p.drawString(50, y, f"TOTAL ITEMS: {total_items}")
        y -= 30
        
        p.setFont("Helvetica-Bold", 10)
        p.drawString(50, y, "Driver Signature: ____________________")
        y -= 25
        p.drawString(50, y, "Date Loaded: ____________________")

        p.showPage()
        p.save()

        buffer.seek(0)
        filename = f"loading_sheet_{sheet_id if sheet_id else 'current'}_{created_at.strftime('%Y%m%d_%H%M')}.pdf"
        
        return Response(
            buffer.getvalue(),
            mimetype='application/pdf',
            headers={"Content-Disposition": f"attachment;filename={filename}"}
        )
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        return f"Error generating PDF: {str(e)}", 500

@app.route('/create-loading-sheet')
def create_loading_sheet():
    """Create a new loading sheet by saving the current one to MongoDB and clearing the session."""
    # Check if there's a current loading sheet in session
    if 'current_loading_sheet' in session:
        current_sheet = session['current_loading_sheet']
        items = current_sheet.get('items', [])
        total_items = current_sheet.get('total_items', 0)
        
        # Parse created_at from session
        created_at_str = current_sheet.get('created_at')
        if isinstance(created_at_str, str):
            try:
                created_at = datetime.fromisoformat(created_at_str)
            except ValueError:
                created_at = datetime.now(NAIROBI_TZ)
        else:
            created_at = datetime.now(NAIROBI_TZ)

        # Generate a unique loading sheet ID
        loading_sheet_id = f"LOAD_{datetime.now(NAIROBI_TZ).strftime('%Y%m%d_%H%M%S')}"
        
        # Save the current sheet to MongoDB
        db["loading_sheets"].insert_one({
            'loading_sheet_id': loading_sheet_id,
            'items': items,
            'total_items': total_items,
            'created_at': created_at,
            'user_id': session['user']['_id']
        })
        
        # Log the action
        log_user_action('Saved Loading Sheet', f"Saved loading sheet {loading_sheet_id} with {total_items} items")
        
        # Clear the current sheet from session
        session.pop('current_loading_sheet')
    
    log_user_action('Created New Loading Sheet', 'Started a fresh loading sheet')
    return redirect(url_for('loading_sheets'))

@app.route('/get_loading_sheet/<sheet_id>')
def get_loading_sheet(sheet_id):
    try:
        sheet = db["loading_sheets"].find_one({'loading_sheet_id': sheet_id})
        if not sheet:
            return jsonify({"error": "Loading sheet not found"}), 404
        
        # Convert ObjectId to string if needed
        if '_id' in sheet:
            sheet['_id'] = str(sheet['_id'])
        
        # Ensure created_at is properly formatted
        created_at_field = sheet.get('created_at')
        if isinstance(created_at_field, str):
            try:
                sheet['created_at'] = datetime.fromisoformat(created_at_field).isoformat()
            except ValueError:
                sheet['created_at'] = datetime.now(NAIROBI_TZ).isoformat()
        else:
            sheet['created_at'] = created_at_field.isoformat()
        
        return jsonify(sheet)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# EXPENSES

UPLOAD_FOLDER = 'static/uploads/receipts'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/expenses', methods=['POST'])
def expenses():
    # 1. Authorization Check
    if session.get('user', {}).get('role') != 'manager':
        return jsonify({'error': 'Unauthorized: Managers only'}), 403

    try:
        # 2. Extract Basic Data
        category = request.form.get('category')
        amount = float(request.form.get('amount', 0))
        description = request.form.get('description', '')
        
        # 3. Initialize the Entry (Prevents "variable not defined" errors)
        expense_entry = {
            'description': description,
            'amount': amount,
            'category': category,
            'metadata': {},
            'receipt_url': None,
            'date': datetime.now(NAIROBI_TZ),
            'user_id': session['user']
        }

        # 4. Apply Category-Specific Logic
        if category == 'Stock':
            stock_name = request.form.get('stock_name')
            unit_price = request.form.get('unit_price')
            expense_entry['metadata'] = {
                'stock_name': stock_name,
                'unit_price': float(unit_price) if unit_price else 0
            }
            # Enhance description for the Audit Trail
            expense_entry['description'] = f"Stock Purchase: {stock_name} - {description}"
            
        elif category in ['Fuel', 'Lunch']:
            expense_entry['description'] = f"{category}: {description}"
            
        elif category == 'Other':
            expense_entry['description'] = f"Other: {description}"

        # 5. Handle the Receipt Image
        if 'receipt' in request.files:
            file = request.files['receipt']
            if file and file.filename != '':
                filename = secure_filename(f"rcpt_{datetime.now().timestamp()}_{file.filename}")
                upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(upload_path)
                expense_entry['receipt_url'] = f"/static/uploads/receipts/{filename}"

        # 6. Database Operations
        db["expenses"].insert_one(expense_entry)
        
        # Log to the Stock/Audit Log
        log_stock_change(category, expense_entry['description'], 'expense', -amount, 1)

        # 7. Return JSON for the Frontend Fetch
        return jsonify({
            'status': 'success',
            'redirect_url': url_for('dashboard')
        }), 200

    except Exception as e:
        print(f"DEBUG ERROR: {str(e)}") # Visible in your terminal
        return jsonify({'error': f"Internal Server Error: {str(e)}"}), 500



# ============================================================================
# COMPARISON API - Compare Previous vs Current Period Performance
# ============================================================================

@app.route('/api/comparison')
@login_required
def get_comparison_data():
    """
    Compare previous period vs current period NET sales
    NET = Total Sales - Expenses (Debt is tracked separately, not subtracted from NET)
    
    Modes:
    - day: Yesterday vs Today
    - week: Last Week vs This Week
    - month: Last Month vs This Month
    
    Returns JSON with previous/current stats and growth metrics
    """
    mode = request.args.get('mode', 'day')
    now = datetime.now(NAIROBI_TZ)
    
    print(f"📊 Comparison API called: mode={mode}")
    
    # ============================================================================
    # STEP 1: CALCULATE DATE RANGES BASED ON MODE
    # ============================================================================
    
    if mode == 'day':
        # Current Period: Today (00:00:00 to 23:59:59)
        current_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        current_end = now.replace(hour=23, minute=59, second=59, microsecond=999999)
        
        # Previous Period: Yesterday (00:00:00 to 23:59:59)
        previous_start = current_start - timedelta(days=1)
        previous_end = previous_start.replace(hour=23, minute=59, second=59, microsecond=999999)
        
        current_label = "Today"
        previous_label = "Yesterday"
    
    elif mode == 'week':
        # Current Period: This Week (Monday 00:00:00 to Sunday 23:59:59)
        days_since_monday = now.weekday()  # Monday=0, Sunday=6
        current_start = now - timedelta(days=days_since_monday)
        current_start = current_start.replace(hour=0, minute=0, second=0, microsecond=0)
        current_end = current_start + timedelta(days=7) - timedelta(seconds=1)
        
        # Previous Period: Last Week (7 days before current week)
        previous_start = current_start - timedelta(days=7)
        previous_end = previous_start + timedelta(days=7) - timedelta(seconds=1)
        
        current_label = "This Week"
        previous_label = "Last Week"
    
    elif mode == 'month':
        # Current Period: This Month (1st 00:00:00 to last day 23:59:59)
        current_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        # Calculate last day of current month
        if now.month == 12:
            next_month_start = datetime(now.year + 1, 1, 1, tzinfo=NAIROBI_TZ)
        else:
            next_month_start = datetime(now.year, now.month + 1, 1, tzinfo=NAIROBI_TZ)
        
        current_end = next_month_start - timedelta(seconds=1)
        
        # Previous Period: Last Month
        previous_end = current_start - timedelta(seconds=1)
        previous_start = previous_end.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        current_label = now.strftime('%B %Y')  # e.g., "January 2026"
        previous_label = previous_start.strftime('%B %Y')  # e.g., "December 2025"
    
    else:
        # Default to day comparison if invalid mode
        current_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        current_end = now.replace(hour=23, minute=59, second=59, microsecond=999999)
        previous_start = current_start - timedelta(days=1)
        previous_end = previous_start.replace(hour=23, minute=59, second=59, microsecond=999999)
        current_label = "Today"
        previous_label = "Yesterday"
    
    print(f"""
    📅 Date Ranges Calculated:
    Previous: {previous_start.strftime('%Y-%m-%d %H:%M')} → {previous_end.strftime('%Y-%m-%d %H:%M')}
    Current:  {current_start.strftime('%Y-%m-%d %H:%M')} → {current_end.strftime('%Y-%m-%d %H:%M')}
    """)
    
    # ============================================================================
    # STEP 2: DEFINE CALCULATION FUNCTION
    # ============================================================================
    
    def calculate_period_stats(start_date, end_date):
        """
        Calculate financial stats for a given period
        
        Returns:
        - total_sales: All money collected (including partial payments on old orders)
        - total_expenses: All expenses incurred
        - total_debt: Debt accumulated from new orders in this period
        - net: Total Sales - Total Expenses
        - orders_count: Number of orders created in this period
        """
        total_sales = 0
        total_expenses = 0
        total_debt = 0
        orders_count = 0
        
        # Fetch all orders once (performance optimization)
        all_orders = list(orders_collection.find())
        
        # ========================================
        # CALCULATE SALES FROM NEW ORDERS
        # ========================================
        for order in all_orders:
            order_date = process_date(order.get('date'))
            
            # Check if order was created in this period
            if order_date and start_date <= order_date <= end_date:
                payment = float(order.get('payment', 0))
                balance = float(order.get('balance', 0))
                
                # Add payment to total sales
                total_sales += payment
                
                # Add unpaid balance to debt
                total_debt += balance
                
                # Count this order
                orders_count += 1
                
                print(f"  📦 Order in period: {order.get('receipt_id')} | Sales: +{payment} | Debt: +{balance}")
        
        # ========================================
        # CALCULATE PARTIAL PAYMENTS ON OLD ORDERS
        # ========================================
        for order in all_orders:
            order_date = process_date(order.get('date'))
            
            # Skip if order was created in this period (already counted above)
            if order_date and start_date <= order_date <= end_date:
                continue
            
            # Check payment history for payments made in this period
            payment_history = order.get('payment_history', [])
            for entry in payment_history:
                pay_date = process_date(entry.get('date'))
                amount = float(entry.get('amount', 0))
                
                # If payment was made in this period, count it
                if pay_date and amount > 0 and start_date <= pay_date <= end_date:
                    total_sales += amount
                    print(f"  💳 Old order payment: {order.get('receipt_id')} | +{amount}")
        
        # ========================================
        # ADD RETAIL COLLECTION (LEGACY)
        # ========================================
        retail_payments = db["retail"].find({"date": {"$gte": start_date, "$lte": end_date}})
        for retail in retail_payments:
            amount = float(retail.get('amount', 0))
            if amount > 0:
                total_sales += amount
                print(f"  🛍️ Retail collection: +{amount}")
        
        # ========================================
        # CALCULATE EXPENSES
        # ========================================
        expenses_in_period = db["expenses"].find({"date": {"$gte": start_date, "$lte": end_date}})
        for expense in expenses_in_period:
            amount = float(expense.get('amount', 0))
            total_expenses += amount
            print(f"  💸 Expense: -{amount}")
        
        # ========================================
        # CALCULATE NET (Sales - Expenses)
        # ========================================
        # Note: We don't subtract debt from NET because debt is money owed, not money lost
        net = total_sales - total_expenses
        
        return {
            'sales': round(total_sales, 2),
            'expenses': round(total_expenses, 2),
            'debt': round(total_debt, 2),
            'net': round(net, 2),
            'orders_count': orders_count
        }
    
    # ============================================================================
    # STEP 3: CALCULATE STATS FOR BOTH PERIODS
    # ============================================================================
    
    print(f"\n📊 Calculating stats for PREVIOUS period ({previous_label})...")
    previous_stats = calculate_period_stats(previous_start, previous_end)
    
    print(f"\n📊 Calculating stats for CURRENT period ({current_label})...")
    current_stats = calculate_period_stats(current_start, current_end)
    
    # ============================================================================
    # STEP 4: CALCULATE GROWTH METRICS
    # ============================================================================
    
    # Absolute growth (difference in NET)
    growth = current_stats['net'] - previous_stats['net']
    
    # Percentage growth (avoid division by zero)
    if previous_stats['net'] > 0:
        growth_percent = (growth / previous_stats['net']) * 100
    elif previous_stats['net'] == 0 and current_stats['net'] > 0:
        growth_percent = 100  # From 0 to positive = 100% growth
    elif previous_stats['net'] == 0 and current_stats['net'] == 0:
        growth_percent = 0  # No change
    else:
        # Previous was negative, current is less negative or positive
        growth_percent = 0  # Edge case, handle as no percentage
    
    
    # ============================================================================
    # STEP 6: RETURN JSON RESPONSE
    # ============================================================================
    
    response_data = {
        'success': True,
        'mode': mode,
        
        # Previous Period Data
        'previous_label': previous_label,
        'previous_sales': previous_stats['sales'],
        'previous_expenses': previous_stats['expenses'],
        'previous_debt': previous_stats['debt'],
        'previous_net': previous_stats['net'],
        'previous_orders_count': previous_stats['orders_count'],
        
        # Current Period Data
        'current_label': current_label,
        'current_sales': current_stats['sales'],
        'current_expenses': current_stats['expenses'],
        'current_debt': current_stats['debt'],
        'current_net': current_stats['net'],
        'current_orders_count': current_stats['orders_count'],
        
        # Growth Metrics
        'growth': round(growth, 2),
        'growth_percent': round(growth_percent, 1),
        'is_positive': growth >= 0
    }
    
    print(f"✅ Comparison data ready, sending response...")
    
    return jsonify(response_data)


# ============================================================================
# HISTORICAL REPORTS ROUTES
# ============================================================================
@app.route('/reports')
def reports():
    time_filter = request.args.get('time', 'month')
    now = datetime.now(NAIROBI_TZ)

    # 1. Precise Date Boundary Setup
    if time_filter == 'day':
        start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'week':
        start = (now - timedelta(days=now.weekday())).replace(hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'month':
        start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'year':
        start = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
    else:
        start = now - timedelta(days=30)

    # 2. Optimized Metrics Aggregation
    
    pipeline_metrics = [
        {
            "$match": {
                "$or": [
                    {"date": {"$gte": start}},
                    {"payment_history.date": {"$gte": start}}
                ]
            }
        },
        {
            "$project": {
                "order_type": 1,
                "is_new_order": {"$gte": ["$date", start]},
                "revenue": {"$add": [{"$ifNull": ["$payment", 0]}, {"$ifNull": ["$balance", 0]}]},
                "balance": {"$ifNull": ["$balance", 0]},
                "period_payments": {
                    "$filter": {
                        "input": {"$ifNull": ["$payment_history", []]},
                        "as": "p",
                        "cond": {"$gte": ["$$p.date", start]}
                    }
                }
            }
        },
        {
            "$group": {
                "_id": "$order_type",
                "total_sales": {"$sum": {"$cond": ["$is_new_order", "$revenue", 0]}},
                "total_debt": {"$sum": {"$cond": ["$is_new_order", "$balance", 0]}},
                "total_paid": {"$sum": {"$sum": "$period_payments.amount"}},
                "mpesa_total": {
                    "$sum": {
                        "$reduce": {
                            "input": "$period_payments",
                            "initialValue": 0,
                            "in": {"$add": ["$$value", {"$cond": [{"$eq": ["$$this.payment_type", "mpesa"]}, "$$this.amount", 0]}]}
                        }
                    }
                }
            }
        }
    ]

    metrics_res = list(orders_collection.aggregate(pipeline_metrics))
    
    # Process Metrics with Guaranteed 0.0 values (Prevents Undefined Error)
    m_data = {
        'wholesale': {'sales':0.0, 'paid':0.0, 'debt':0.0, 'bank':0.0}, 
        'retail': {'sales':0.0, 'paid':0.0, 'debt':0.0, 'bank':0.0}
    }
    
    for m in metrics_res:
        m_type = m['_id'] if m['_id'] in m_data else 'wholesale'
        m_data[m_type] = {
            'sales': float(m.get('total_sales', 0)), 
            'paid': float(m.get('total_paid', 0)), 
            'debt': float(m.get('total_debt', 0)), 
            'bank': float(m.get('mpesa_total', 0))
        }

    # 3. Aggregation for Cumulative Growth Chart
    pipeline_timeline = [
        {"$unwind": "$payment_history"},
        {"$match": {"payment_history.date": {"$gte": start}}},
        {
            "$group": {
                "_id": {"$dateToString": {"format": "%Y-%m-%d", "date": "$payment_history.date"}},
                "daily_sum": {"$sum": "$payment_history.amount"}
            }
        },
        {"$sort": {"_id": 1}}
    ]
    
    timeline_res = list(orders_collection.aggregate(pipeline_timeline))
    
    labels, cumulative_data, running_total = [], [], 0
    for day in timeline_res:
        labels.append(day['_id'])
        running_total += float(day.get('daily_sum', 0))
        cumulative_data.append(running_total)

    # 4. Final Chart Object Construction
    chart_data = {
        'sales_vs_debts': {
            'data': [m_data['retail']['sales'], m_data['wholesale']['sales'], 
                     m_data['retail']['debt'], m_data['wholesale']['debt']]
        },
        'paid_vs_debt': {
            'data': [m_data['retail']['paid'], m_data['wholesale']['paid'], 
                     m_data['retail']['debt'], m_data['wholesale']['debt']]
        },
        'money_in_bank': {
            'data': [m_data['retail']['bank'], m_data['wholesale']['bank']]
        },
        'timeline_labels': labels if labels else [now.strftime('%Y-%m-%d')],
        'cumulative_bank_data': cumulative_data if cumulative_data else [0]
    }

    # 5. Fetch Activity Logs (Limited for performance)
    orders = list(orders_collection.find({"date": {"$gte": start}}).sort("date", -1).limit(50))
    for o in orders:
        o['date'] = parse_mongo_date(o.get('date'))
        # Ensure total_amount is a number for the |round filter
        o['total_amount'] = float(o.get('payment', 0) + o.get('balance', 0))
        o['receipt_id'] = o.get('receipt_id', str(o['_id']))

    # Final safe calculation for total_debt
    total_debt_value = m_data['retail']['debt'] + m_data['wholesale']['debt']

    recent_reports = list(db['sales_reports'].find().sort('generated_at', -1).limit(6))

    return render_template('reports.html', 
                         chart_data=chart_data, 
                         time_filter=time_filter, 
                         total_debt=total_debt_value,
                         recent_reports=recent_reports,
                         orders=orders)


@app.route('/sales_report')
@login_required
def sales_report():
    """
    Dynamic sales report - handles daily/weekly/monthly/yearly
    Displays printable receipt-style summary
    """
    # Get time filter from query params (default: day)
    time_filter = request.args.get('time', 'day')
    now = datetime.now(NAIROBI_TZ)
    
    # ============================================================================
    # CALCULATE DATE RANGE BASED ON TIME FILTER
    # ============================================================================
    if time_filter == 'day':
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = now.replace(hour=23, minute=59, second=59, microsecond=999999)
        period_label = now.strftime('%d %B %Y')
        report_title = "DAILY SALES REPORT"
    
    elif time_filter == 'week':
        # Start of week (Monday)
        start_date = now - timedelta(days=now.weekday())
        start_date = start_date.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date + timedelta(days=7)
        period_label = f"{start_date.strftime('%d %b')} - {(end_date - timedelta(days=1)).strftime('%d %b %Y')}"
        report_title = "WEEKLY SALES REPORT"
    
    elif time_filter == 'month':
        # Start of month
        start_date = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        # End of month
        next_month = start_date + timedelta(days=32)
        end_date = next_month.replace(day=1) - timedelta(seconds=1)
        period_label = now.strftime('%B %Y')
        report_title = "MONTHLY SALES REPORT"
    
    elif time_filter == 'year':
        # Start of year
        start_date = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
        end_date = start_date.replace(year=start_date.year + 1) - timedelta(seconds=1)
        period_label = str(now.year)
        report_title = "YEARLY SALES REPORT"
    
    else:
        # Default to today
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = now.replace(hour=23, minute=59, second=59, microsecond=999999)
        period_label = now.strftime('%d %B %Y')
        report_title = "DAILY SALES REPORT"
    
    # ============================================================================
    # CALCULATE SALES STATISTICS
    # ============================================================================
    total_wholesale_revenue = 0
    total_retail_revenue = 0
    total_wholesale_paid = 0
    total_retail_paid = 0
    total_debt = 0
    total_mpesa = 0
    total_cash = 0
    period_orders = []
    
    all_orders = list(orders_collection.find())
    
    # === 1. ORDERS CREATED IN THIS PERIOD ===
    for order in all_orders:
        order_type = order.get('order_type', 'wholesale').lower()
        payment = float(order.get('payment', 0))
        balance = float(order.get('balance', 0))
        receipt_id = order.get('receipt_id', str(order['_id']))
        payment_breakdown = order.get('payment_breakdown')
        payment_type = order.get('payment_type', '').lower().strip()
        payment_history = order.get('payment_history', [])
        order_date = process_date(order.get('date'))
        
        # Check if order is in this period
        is_in_period = order_date and start_date <= order_date <= end_date
        
        if is_in_period:
            # Revenue = payment + balance
            revenue = payment + balance
            if order_type == 'wholesale':
                total_wholesale_revenue += revenue
            else:
                total_retail_revenue += revenue
            
            period_orders.append({
                'receipt_id': receipt_id,
                'order_type': order_type,
                'payment': payment,
                'balance': balance
            })
            
            # Payments made on this order
            if payment > 0:
                if order_type == 'wholesale':
                    total_wholesale_paid += payment
                else:
                    total_retail_paid += payment
                
                total_cash_for_order = 0
                total_mpesa_for_order = 0
                
                # Check payment history
                if payment_history:
                    for entry in payment_history:
                        entry_date = process_date(entry.get('date'))
                        if entry_date and start_date <= entry_date <= end_date:
                            entry_amount = float(entry.get('amount', 0))
                            entry_type = entry.get('payment_type', 'cash').lower()
                            
                            if entry_type == 'mpesa':
                                total_mpesa_for_order += entry_amount
                            else:
                                total_cash_for_order += entry_amount
                
                # Fallback to breakdown
                elif payment_breakdown:
                    total_mpesa_for_order = float(payment_breakdown.get('mpesa', 0))
                    total_cash_for_order = float(payment_breakdown.get('cash', 0))
                
                # Final fallback
                elif payment_type == 'mpesa':
                    total_mpesa_for_order = payment
                else:
                    total_cash_for_order = payment
                
                total_mpesa += total_mpesa_for_order
                total_cash += total_cash_for_order
                print(f"{time_filter.upper()} {order_type.upper()}: {receipt_id} | MPESA: {total_mpesa_for_order} | CASH: {total_cash_for_order}")
            
            # Debt from this order
            if balance > 0:
                total_debt += balance
    
    # === 2. PARTIAL PAYMENTS ON OLD ORDERS (IN THIS PERIOD) ===
    for order in all_orders:
        order_type = order.get('order_type', 'wholesale').lower()
        receipt_id = order.get('receipt_id', str(order['_id']))
        payment_history = order.get('payment_history', [])
        order_date = process_date(order.get('date'))
        
        # Skip if order was created in this period
        if order_date and start_date <= order_date <= end_date:
            continue
        
        for entry in payment_history:
            pay_date = process_date(entry.get('date'))
            amount = float(entry.get('amount', 0))
            entry_payment_type = entry.get('payment_type', 'cash').lower()
            entry_breakdown = entry.get('payment_breakdown')
            
            if pay_date and amount > 0 and start_date <= pay_date <= end_date:
                if order_type == 'wholesale':
                    total_wholesale_paid += amount
                else:
                    total_retail_paid += amount
                
                if entry_breakdown:
                    entry_mpesa = float(entry_breakdown.get('mpesa', 0))
                    entry_cash = float(entry_breakdown.get('cash', 0))
                    total_mpesa += entry_mpesa
                    total_cash += entry_cash
                    print(f"OLD {order_type.upper()} PAID IN PERIOD (DUAL): {receipt_id} | MPESA: {entry_mpesa} | CASH: {entry_cash}")
                elif entry_payment_type == 'mpesa':
                    total_mpesa += amount
                    print(f"OLD {order_type.upper()} PAID IN PERIOD (MPESA): {receipt_id} | +{amount}")
                else:
                    total_cash += amount
                    print(f"OLD {order_type.upper()} PAID IN PERIOD (CASH): {receipt_id} | +{amount}")
    
    # === 3. RETAIL COLLECTION (LEGACY) ===
    for retail in db["retail"].find({"date": {"$gte": start_date, "$lt": end_date}}):
        amount = float(retail.get('amount', 0))
        if amount > 0:
            total_retail_paid += amount
            total_cash += amount
            total_retail_revenue += amount
    
    # === 4. EXPENSES ===
    total_expenses = sum(
        float(e.get('amount', 0))
        for e in db["expenses"].find({"date": {"$gte": start_date, "$lt": end_date}})
    )
    
    # === FINAL CALCULATIONS ===
    total_sales = total_wholesale_paid + total_retail_paid
    net = total_sales - total_expenses
    
    # === CONSOLE PRINT SUMMARY ===
    print(f"""
    ========================================
    {report_title} - {period_label}
    ========================================
    Revenue (Wholesale): KSh {total_wholesale_revenue:,.2f}
    Revenue (Retail):    KSh {total_retail_revenue:,.2f}
    
    Paid (Wholesale):    KSh {total_wholesale_paid:,.2f}
    Paid (Retail):       KSh {total_retail_paid:,.2f}
    
    Payment Breakdown:
    - Cash:              KSh {total_cash:,.2f}
    - M-Pesa:            KSh {total_mpesa:,.2f}
    
    Total Sales:         KSh {total_sales:,.2f}
    Total Expenses:      KSh {total_expenses:,.2f}
    Outstanding Debt:    KSh {total_debt:,.2f}
    
    NET PROFIT:          KSh {net:,.2f}
    ========================================
    Orders Processed: {len(period_orders)}
    ========================================
    """)
    report_record = {
        'report_type': time_filter,
        'period_start': start_date,
        'period_end': end_date,
        'period_label': period_label,
        'generated_at': now,
        'generated_by': session.get('user', {}).get('email', 'unknown'),
        'stats': {
            'total_wholesale_revenue': total_wholesale_revenue,
            'total_retail_revenue': total_retail_revenue,
            'total_wholesale_paid': total_wholesale_paid,
            'total_retail_paid': total_retail_paid,
            'total_debt': total_debt,
            'total_expenses': total_expenses,
            'total_sales': total_sales,
            'net': net,
            'orders_count': len(period_orders),
            'total_mpesa': total_mpesa,
            'total_cash': total_cash
        },
        'orders': period_orders  # Optional: Remove if too large
    }
        
    db['sales_reports'].insert_one(report_record)
    #print (f"💾 Report saved: {period_label}")
    # ============================================================================
    # RENDER TEMPLATE
    # ============================================================================
    return render_template('daily_sales_report.html',
        report_title=report_title,
        date=period_label,
        time_generated=now.strftime('%H:%M:%S'),
        time_filter=time_filter,
        total_wholesale_revenue=total_wholesale_revenue,
        total_retail_revenue=total_retail_revenue,
        total_wholesale_paid=total_wholesale_paid,
        total_retail_paid=total_retail_paid,
        total_debt=total_debt,
        total_expenses=total_expenses,
        total_sales=total_sales,
        net=net,
        orders_count=len(period_orders),
        today_orders=period_orders,
        total_mpesa=total_mpesa,
        total_cash=total_cash
    )

@app.route('/sales_reports_history')
@login_required
def sales_reports_history():
    """View all past generated reports"""
    console.log("📂 Loading sales reports history")
    
    reports = list(db['sales_reports'].find().sort('generated_at', -1).limit(50))
    
    return render_template('sales_reports_history.html', reports=reports)


@app.route('/view_report/<report_id>')
@login_required
def view_report(report_id):
    """View/Print a saved report"""
    from bson import ObjectId
    
    console.log(f"👁️ Viewing report: {report_id}")
    
    try:
        report = db['sales_reports'].find_one({'_id': ObjectId(report_id)})
    except:
        flash('Invalid report ID', 'error')
        return redirect(url_for('sales_reports_history'))
    
    if not report:
        flash('Report not found', 'error')
        return redirect(url_for('sales_reports_history'))
    
    # Extract saved stats
    stats = report.get('stats', {})
    
    return render_template('daily_sales_report.html',
        report_title=f"{report['report_type'].upper()} SALES REPORT",
        date=report['period_label'],
        time_generated=report['generated_at'].strftime('%H:%M:%S'),
        total_wholesale_revenue=stats.get('total_wholesale_revenue', 0),
        total_retail_revenue=stats.get('total_retail_revenue', 0),
        total_wholesale_paid=stats.get('total_wholesale_paid', 0),
        total_retail_paid=stats.get('total_retail_paid', 0),
        total_debt=stats.get('total_debt', 0),
        total_expenses=stats.get('total_expenses', 0),
        total_sales=stats.get('total_sales', 0),
        net=stats.get('net', 0),
        orders_count=stats.get('orders_count', 0),
        total_mpesa=stats.get('total_mpesa', 0),
        total_cash=stats.get('total_cash', 0),
        today_orders=report.get('orders', [])
    )


@app.route('/api/delete_report/<report_id>', methods=['DELETE'])
@login_required
def delete_report(report_id):
    """Delete a saved report"""
    from bson import ObjectId
    
    console.log(f"🗑️ Deleting report: {report_id}")
    
    try:
        result = db['sales_reports'].delete_one({'_id': ObjectId(report_id)})
        
        if result.deleted_count > 0:
            console.log("✅ Report deleted successfully")
            return jsonify({'success': True, 'message': 'Report deleted'})
        else:
            console.log("⚠️ Report not found")
            return jsonify({'success': False, 'message': 'Report not found'}), 404
    
    except Exception as e:
        console.log(f"❌ Error deleting report: {str(e)}")
        return jsonify({'success': False, 'message': str(e)}), 500

# ============================================================================
# EXPORT ROUTE (Future Implementation)
# ============================================================================
@app.route('/export_sales_report/<format>')
@login_required
def export_sales_report(format):
    """
    Export sales report as PDF or DOCX
    TODO: Implement PDF and DOCX generation
    """
    time_filter = request.args.get('time', 'day')
    
    # TODO: Reuse calculation logic from sales_report() route
    # TODO: Generate PDF using reportlab or weasyprint
    # TODO: Generate DOCX using python-docx
    
    flash('Export feature coming soon!', 'info')
    return redirect(url_for('sales_report', time=time_filter))

@app.route('/export_report')
def export_report():
    report_type = request.args.get('type')
    time_filter = request.args.get('time', 'month')
    now = datetime.now(NAIROBI_TZ)

    if report_type == 'daily_sales':
        return redirect(url_for('daily_sales_report'))

    # Set time range
    if time_filter == 'day':
        start = now.replace(hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'week':
        start = now - timedelta(days=now.weekday())
        start = start.replace(hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'month':
        start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'year':
        start = now.replace(month=1, day=1, hour=0, minute=0, second=0, microsecond=0)
    else:
        start = None

    # Generate PDF
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    # Header
    p.setFont("Helvetica-Bold", 16)
    p.drawCentredString(width / 2, height - 40, "Dreamland Distributors")
    p.setFont("Helvetica", 10)
    p.drawCentredString(width / 2, height - 60, "P.O Box 123-00200 Nairobi | Phone: 0725 530632 | Email: info@dreamland.co.ke")
    p.setFont("Helvetica-Oblique", 8)
    p.drawCentredString(width / 2, height - 75, "Financial Report")
    p.line(40, height - 85, width - 40, height - 85)

    # Report Title and Metadata
    p.setFont("Helvetica-Bold", 12)
    report_title = f"{report_type.replace('_', ' ').title()} Report - {time_filter.capitalize()}"
    p.drawCentredString(width / 2, height - 110, report_title)
    p.setFont("Helvetica", 9)
    p.drawString(40, height - 130, f"Generated on: {now.strftime('%d/%m/%Y %H:%M')}")
    p.drawString(40, height - 145, f"Generated by: {session['user']['first_name']} {session['user']['last_name']}")
    p.drawString(width - 150, height - 130, f"Period: {time_filter.capitalize()}")

    y = height - 170

    if report_type == 'stock':
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, "Product")
        p.drawString(170, y, "Category")
        p.drawString(320, y, "Quantity")
        p.drawString(420, y, "Value (KES)")
        p.drawString(510, y, "Date")
        y -= 10
        p.line(40, y, width - 40, y)
        y -= 10
        
        stock_query = {"timestamp": {"$gte": start}} if start else {}
        stock_logs = list(stock_logs_collection.find(stock_query).sort("timestamp", -1))
        p.setFont("Helvetica", 9)
        total_movement = 0
        total_value = 0
        
        for log in stock_logs:
            timestamp = process_date(log.get('timestamp'))
            
            if y < 60:
                p.showPage()
                p.setFont("Helvetica-Bold", 10)
                p.drawString(40, height - 50, "Product")
                p.drawString(170, height - 50, "Category")
                p.drawString(320, height - 50, "Quantity")
                p.drawString(420, height - 50, "Value (KES)")
                p.drawString(510, height - 50, "Date")
                p.line(40, height - 60, width - 40, height - 60)
                p.setFont("Helvetica", 9)
                y = height - 80
                
            qty = log.get('quantity', 0)
            price = log.get('price_per_unit', 0)
            value = qty * price
            
            p.drawString(40, y, log.get('subtype', 'Unknown'))
            p.drawString(170, y, log.get('product_type', 'Unknown'))
            p.drawString(320, y, str(qty))
            p.drawString(420, y, f"{value:.2f}")
            p.drawString(510, y, timestamp.strftime('%d/%m/%Y'))
            
            total_movement += qty
            total_value += value
            y -= 15

        y -= 10
        p.line(40, y, width - 40, y)
        y -= 15
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, f"Total Items Moved: {total_movement}")
        p.drawString(320, y, f"Total Value: {total_value:.2f} KES")

    elif report_type == 'user':
        user_query = {"date": {"$gte": start}} if start else {}
        orders = list(orders_collection.find(user_query).sort("date", -1))
        user_data = {}
        
        for order in orders:
            order_date = process_date(order.get('date'))
            salesperson = order.get('salesperson_name', 'Unknown')
            if salesperson not in user_data:
                user_data[salesperson] = {
                    'orders': [],
                    'total_debt': 0,
                    'total_sales': 0,
                    'total_items': 0
                }
            user_data[salesperson]['orders'].append(order)
            user_data[salesperson]['total_debt'] += order.get('balance', 0)
            user_data[salesperson]['total_sales'] += order.get('payment', 0)
            user_data[salesperson]['total_items'] += process_items(order.get('items', []))

        p.setFont("Helvetica", 9)
        for salesperson, data in user_data.items():
            if y < 60:
                p.showPage()
                y = height - 50

            p.setFont("Helvetica-Bold", 11)
            p.drawString(40, y, f"User: {salesperson}")
            y -= 20

            p.setFont("Helvetica-Bold", 9)
            p.drawString(40, y, "Order ID")
            p.drawString(120, y, "Shop")
            p.drawString(220, y, "Items Sold")
            p.drawString(300, y, "Debt (KES)")
            p.drawString(380, y, "Sales (KES)")
            p.drawString(460, y, "Date")
            y -= 10
            p.line(40, y, width - 40, y)
            y -= 10
            
            p.setFont("Helvetica", 9)
            for order in data['orders']:
                if y < 60:
                    p.showPage()
                    p.setFont("Helvetica-Bold", 9)
                    p.drawString(40, height - 50, "Order ID")
                    p.drawString(120, height - 50, "Shop")
                    p.drawString(220, height - 50, "Items Sold")
                    p.drawString(300, height - 50, "Debt (KES)")
                    p.drawString(380, height - 50, "Sales (KES)")
                    p.drawString(460, height - 50, "Date")
                    p.line(40, height - 60, width - 40, height - 60)
                    p.setFont("Helvetica", 9)
                    y = height - 80
                    
                p.drawString(40, y, order.get('receipt_id', str(order['_id'])))
                p.drawString(120, y, order.get('shop_name', 'Unknown'))
                p.drawString(220, y, str(process_items(order.get('items', []))))
                p.drawString(300, y, f"{order.get('balance', 0):.2f}")
                p.drawString(380, y, f"{order.get('payment', 0):.2f}")
                p.drawString(460, y, process_date(order.get('date')).strftime('%d/%m/%Y'))
                y -= 15

            y -= 5
            p.line(40, y, width - 40, y)
            y -= 15
            p.setFont("Helvetica-Bold", 10)
            p.drawString(40, y, f"Summary for {salesperson}:")
            p.drawString(220, y, f"Orders: {len(data['orders'])}")
            p.drawString(300, y, f"Items: {data['total_items']}")
            p.drawString(380, y, f"Debt: {data['total_debt']:.2f} KES")
            p.drawString(460, y, f"Sales: {data['total_sales']:.2f} KES")
            y -= 25

    elif report_type == 'debt':
        debt_query = {"balance": {"$gt": 0}, "date": {"$gte": start}} if start else {"balance": {"$gt": 0}}
        orders = list(orders_collection.find(debt_query).sort("date", -1))
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, "Order ID")
        p.drawString(120, y, "Shop")
        p.drawString(220, y, "Salesperson")
        p.drawString(320, y, "Debt Amount (KES)")
        p.drawString(420, y, "Date")
        y -= 10
        p.line(40, y, width - 40, y)
        y -= 10
        
        p.setFont("Helvetica", 9)
        total_debt = 0
        
        for order in orders:
            order_date = process_date(order.get('date'))
            debt = order.get('balance', 0)
            total_debt += debt
            
            if y < 60:
                p.showPage()
                p.setFont("Helvetica-Bold", 10)
                p.drawString(40, height - 50, "Order ID")
                p.drawString(120, height - 50, "Shop")
                p.drawString(220, height - 50, "Salesperson")
                p.drawString(320, height - 50, "Debt Amount (KES)")
                p.drawString(420, height - 50, "Date")
                p.line(40, height - 60, width - 40, height - 60)
                p.setFont("Helvetica", 9)
                y = height - 80
                
            p.drawString(40, y, order.get('receipt_id', str(order['_id'])))
            p.drawString(120, y, order.get('shop_name', 'Unknown'))
            p.drawString(220, y, order.get('salesperson_name', 'Unknown'))
            p.drawString(320, y, f"{debt:.2f}")
            p.drawString(420, y, order_date.strftime('%d/%m/%Y'))
            y -= 15

        y -= 10
        p.line(40, y, width - 40, y)
        y -= 15
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, f"Total Outstanding Debt: {total_debt:.2f} KES")

    elif report_type == 'sales':
        sales_query = {"date": {"$gte": start}} if start else {}
        orders = list(orders_collection.find(sales_query).sort("date", -1))
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, "Order ID")
        p.drawString(110, y, "Shop")
        p.drawString(220, y, "Salesperson")
        p.drawString(320, y, "Items")
        p.drawString(370, y, "Payment (KES)")
        p.drawString(450, y, "Debt (KES)")
        p.drawString(520, y, "Date")
        y -= 10
        p.line(40, y, width - 40, y)
        y -= 10
        
        p.setFont("Helvetica", 9)
        total_sales = 0
        total_debt = 0
        
        for order in orders:
            order_date = process_date(order.get('date'))
            payment = order.get('payment', 0)
            debt = order.get('balance', 0)
            total_sales += payment
            total_debt += debt
            
            if y < 60:
                p.showPage()
                p.setFont("Helvetica-Bold", 10)
                p.drawString(40, height - 50, "Order ID")
                p.drawString(110, height - 50, "Shop")
                p.drawString(220, height - 50, "Salesperson")
                p.drawString(320, height - 50, "Items")
                p.drawString(370, height - 50, "Payment (KES)")
                p.drawString(450, height - 50, "Debt (KES)")
                p.drawString(520, height - 50, "Date")
                p.line(40, height - 60, width - 40, height - 60)
                p.setFont("Helvetica", 9)
                y = height - 80
                
            p.drawString(40, y, order.get('receipt_id', str(order['_id'])))
            p.drawString(110, y, order.get('shop_name', 'Unknown'))
            p.drawString(220, y, order.get('salesperson_name', 'Unknown'))
            p.drawString(320, y, str(process_items(order.get('items', []))))
            p.drawString(370, y, f"{payment:.2f}")
            p.drawString(450, y, f"{debt:.2f}")
            p.drawString(520, y, order_date.strftime('%d/%m/%Y'))
            y -= 15

        y -= 10
        p.line(40, y, width - 40, y)
        y -= 15
        p.setFont("Helvetica-Bold", 10)
        p.drawString(40, y, f"Total Sales: {total_sales:.2f} KES")
        p.drawString(300, y, f"Total Outstanding Debt: {total_debt:.2f} KES")

    p.setFont("Helvetica-Oblique", 8)
    p.drawString(40, 40, "Dreamland Distributors System © 2025")
    p.drawString(width - 150, 40, f"Page {p.getPageNumber()}")

    p.showPage()
    p.save()
    buffer.seek(0)

    filename = f"{report_type}_report_{now.strftime('%Y%m%d_%H%M')}.pdf"
    return Response(
        buffer,
        mimetype='application/pdf',
        headers={"Content-Disposition": f"attachment;filename={filename}"}
    )


# 1. NOTIFICATIONS ENDPOINT
@app.route('/api/notifications', methods=['GET'])
def get_notifications():
    """Get all notifications or filter by status"""
    try:
        filter_type = request.args.get('filter', 'all')
        view_type = request.args.get('view', 'recent')
        
        query = {}
        now = datetime.now(NAIROBI_TZ)
        
        if filter_type == 'today':
            today_start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            query['created_at'] = {'$gte': today_start}
        elif filter_type == 'week':
            week_ago = now - timedelta(days=7)
            query['created_at'] = {'$gte': week_ago}
        elif filter_type == 'unread':
            query['read'] = False
        
        notifications = list(notifications_collection.find(query).sort('created_at', -1))
        
        if view_type == 'recent':
            notifications = [n for n in notifications if not n.get('read', False)][:5]
        
        result = []
        for notif in notifications:
            if notif.get('type') == 'expiry':
                result.append({
                    'id': notif['notification_id'],
                    'message': notif['message'],
                    'time': notif['created_at'].isoformat(),
                    'unread': not notif.get('read', False),
                    'type': 'expiry',
                    'stock_name': notif.get('stock_name'),
                    'days_left': notif.get('days_left'),
                    'category': 'stock'
                })
            elif notif.get('type') == 'stock_change':
                result.append({
                    'id': notif.get('notification_id'),
                    'message': notif.get('message'),
                    'time': notif.get('created_at').isoformat(),
                    'unread': not notif.get('read', False),
                    'type': 'stock_change',
                    'category': 'stock',
                    'stock_name': notif.get('stock_name'),
                    'action': notif.get('action'),
                    'creator': notif.get('user'),
                })
            else:
                result.append({
                    'id': notif.get('notification_id'),
                    'receiptId': notif.get('receipt_id'),
                    'creator': notif.get('salesperson_name'),
                    'time': notif.get('created_at').isoformat(),
                    'status': notif.get('status'),
                    'shopName': notif.get('shop_name'),
                    'amount': notif.get('amount'),
                    'orderType': notif.get('order_type'),
                    'unread': not notif.get('read', False),
                    'type': notif.get('type'),
                    'category': notif.get('category', 'orders')
                })
        
        unread_count = notifications_collection.count_documents({'read': False})
        
        return jsonify({
            'notifications': result,
            'unread_count': unread_count
        }), 200
        
    except Exception as e:
        logger.error(f"Error fetching notifications: {e}")
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/notifications/<notification_id>/mark-read', methods=['POST'])
def mark_notification_read(notification_id):
    """Mark a notification as read - DOES NOT DELETE"""
    try:
        notif = notifications_collection.find_one({'notification_id': notification_id})
        if not notif:
            return jsonify({'error': 'Notification not found'}), 404
        
        # Just toggle read status, never delete
        new_read_status = not notif.get('read', False)
        
        notifications_collection.update_one(
            {'notification_id': notification_id},
            {'$set': {'read': new_read_status}}
        )
        
        logger.info(f"Notification {notification_id} marked as {'read' if new_read_status else 'unread'}")
        
        return jsonify({'message': 'Notification updated', 'read': new_read_status}), 200
        
    except Exception as e:
        logger.error(f"Error updating notification: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/notifications/<notification_id>/clear', methods=['POST'])
def clear_notification(notification_id):
    """Mark notification as read (clear from dropdown) - DOES NOT DELETE"""
    try:
        result = notifications_collection.update_one(
            {'notification_id': notification_id},
            {'$set': {'read': True}}  # Just mark as read, don't delete
        )
        
        if result.matched_count == 0:
            return jsonify({'error': 'Notification not found'}), 404
        
        logger.info(f"Notification {notification_id} cleared (marked as read)")
        
        return jsonify({'message': 'Notification cleared'}), 200
        
    except Exception as e:
        logger.error(f"Error clearing notification: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/logs')
def get_logs():
    tab = request.args.get('tab', 'orders')
    page = int(request.args.get('page', 1))
    time_filter = request.args.get('time', 'day')
    search_query = request.args.get('search', '').strip()
    per_page = 20
    skip = (page - 1) * per_page

    # CRITICAL: These must match your db["collection_name"] in the functions above
    collection_map = {
        'orders': 'orders',
        'expenses': 'expenses',
        'stock': 'inventory_logs',    
        'user-actions': 'audit_trail' 
    }
    
    collection_name = collection_map.get(tab, 'orders')
    
    # Date Filtering Logic
    now = datetime.now(NAIROBI_TZ)
    if time_filter == 'day':
        start_date = now.replace(hour=0, minute=0, second=0, microsecond=0)
    elif time_filter == 'week':
        start_date = now - timedelta(days=now.weekday())
    elif time_filter == 'month':
        start_date = now.replace(day=1)
    else:
        start_date = now.replace(month=1, day=1)

    query = {'date': {'$gte': start_date}}

    if search_query:
        query['$or'] = [
            {'item': {'$regex': search_query, '$options': 'i'}},
            {'description': {'$regex': search_query, '$options': 'i'}},
            {'details': {'$regex': search_query, '$options': 'i'}}
        ]

    try:
        # Fetch per_page + 1 to check for "Load More"
        raw_data = list(db[collection_name].find(query)
                       .sort('date', -1)
                       .skip(skip)
                       .limit(per_page + 1))

        has_more = len(raw_data) > per_page
        items = raw_data[:per_page]

        # Render the specific partial for the tab
        # Filenames: log_orders.html, log_stock.html, log_expenses.html, log_user-actions.html
        html_content = render_template(f'partials/log_{tab}.html', items=items)

        return jsonify({
            'html': html_content,
            'has_more': has_more
        })
    except Exception as e:
        print(f"API Error: {str(e)}")
        return jsonify({'error': str(e), 'html': ''}), 500

# INIT 0

@app.route('/logout')
def logout():
    """
    Enhanced logout with comprehensive session tracking.
    
    Handles:
    - Manual logout (user clicks button)
    - Abnormal termination detection (power loss, crashes)
    - Session duration tracking
    - Clean session cleanup
    """
    current_user = get_current_user()
    
    if current_user:
        user_email = current_user.get('email', 'unknown')
        login_time_str = session.get('login_time')
        now = datetime.now(NAIROBI_TZ)
        
        # Calculate session duration
        session_duration = None
        if login_time_str:
            try:
                login_time = datetime.fromisoformat(login_time_str)
                session_duration = (now - login_time).total_seconds() / 60  # minutes
            except:
                pass
        
        # Log manual logout
        db['session_logs'].insert_one({
            'email': user_email,
            'action': 'logout',
            'logout_type': 'manual',  # User explicitly clicked logout
            'session_duration_minutes': session_duration,
            'timestamp': now,
            'ip_address': request.remote_addr,
            'user_agent': request.headers.get('User-Agent')
        })
        
        session.clear()
        flash('You have been logged out successfully.', 'success')
    
    return redirect(url_for('splash'))


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5001, debug=True)
